"""
Captures : numéro de figure tenu par Word, repères à glisser sur l'image.

Deux contrats sont vérifiés ici. Le numéro d'une légende est un champ : Word le
recalcule, et le voir changer ne veut pas dire que quelqu'un a écrit dedans. Un
repère, lui, ne porte que son numéro : seule sa position dit qu'on l'a déplacé
— et une fois déplacé, il reste où on l'a mis.
"""

import copy
import os
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn

from src import console
from src.config import DocConfig
from src.generators.word import generate_word_documentation

ANCHOR = qn("wp:anchor")
OFFSET = qn("wp:posOffset")
FIELD_CHAR = qn("w:fldChar")
INSTR_TEXT = qn("w:instrText")
TEXT = qn("w:t")

PLAN = {
    "document": {"template": "", "cover": {}, "properties": {}},
    "styles": {"normal": "Normal", "caption": "Normal", "image": "Normal", "fallback": "Normal"},
    "rendering": {
        "page_break_before_heading_1": False,
        "links": {"enabled": False},
        "image_placeholder": {
            "text_format": "[IMAGE] {description}",
            "caption_format": "Figure {n} — {description}",
            "show_caption": True,
            "numbering": "auto",
            "empty_paragraph_after": False,
        },
    },
    "sections": [
        {
            "id": "visuel",
            "title": "Visuel",
            "level": 1,
            "blocks": [
                {
                    "type": "image",
                    "id": "capture",
                    "description": "Capture du visuel",
                    "markers": {"over": "inputs.champs", "item": "champ", "label": "{{ champ }}"},
                }
            ],
        }
    ],
}


class Harness(unittest.TestCase):
    def setUp(self):
        self._directory = tempfile.TemporaryDirectory()
        self.addCleanup(self._directory.cleanup)
        self.path = os.path.join(self._directory.name, "doc.docx")
        self.template = os.path.join(self._directory.name, "template.docx")
        Document().save(self.template)

    def generate(
        self, rendering: dict | None = None, block: dict | None = None, champs=("1", "2", "3")
    ):
        plan = copy.deepcopy(PLAN)
        plan["document"]["template"] = self.template
        if rendering:
            plan["rendering"]["image_placeholder"].update(rendering)
        if block is not None:
            plan["sections"][0]["blocks"][0].update(block)
        context = {
            "model": None,
            "report": None,
            "inputs": {"champs": list(champs)},
            "styles": plan["styles"],
        }
        with console.silenced():
            generate_word_documentation(DocConfig(plan), context, self.path)

    # ── Lecture du document produit ───────────────────────────────
    def document(self):
        return Document(self.path)

    def anchors(self) -> list:
        return list(self.document().element.body.iter(ANCHOR))

    def offsets(self) -> list[str]:
        return [offset.text for anchor in self.anchors() for offset in anchor.iter(OFFSET)]

    def caption(self):
        return next(p for p in self.document().paragraphs if p.text.startswith("Figure"))

    def texts(self) -> list[str]:
        return [
            p.text
            for p in self.document().paragraphs
            if p.text.strip() and not p.text.startswith("pbi::")
        ]


class NumeroDeFigureTest(Harness):
    def test_le_numero_est_un_champ_word(self):
        self.generate()
        instructions = [t.text for t in self.caption()._p.iter(INSTR_TEXT)]
        self.assertEqual(instructions, [" SEQ Figure \\* ARABIC "])

    def test_le_champ_est_marque_a_recalculer(self):
        """`w:dirty` : Word renumérote de lui-même à l'ouverture du document."""
        self.generate()
        dirty = [c.get(qn("w:dirty")) for c in self.caption()._p.iter(FIELD_CHAR)]
        self.assertIn("true", dirty)

    def test_le_numero_compte_reste_lisible_sans_recalcul(self):
        self.generate()
        self.assertEqual(self.caption().text, "Figure 1 — Capture du visuel")

    def test_numerotation_figee_sur_demande(self):
        self.generate({"numbering": "fixed"})
        self.assertEqual(self.caption().text, "Figure 1 — Capture du visuel")
        self.assertEqual([t.text for t in self.caption()._p.iter(INSTR_TEXT)], [])

    def test_numerotation_desactivee(self):
        self.generate({"numbering": False})
        self.assertEqual(self.caption().text, "Figure  — Capture du visuel")

    def test_ancienne_forme_booleenne_toujours_admise(self):
        self.generate({"numbering": True})
        self.assertEqual(
            [t.text for t in self.caption()._p.iter(INSTR_TEXT)], [" SEQ Figure \\* ARABIC "]
        )


class RepresSurLaCaptureTest(Harness):
    def test_un_repere_par_element(self):
        self.generate(champs=("1", "2", "3", "4"))
        self.assertEqual(len(self.anchors()), 4)

    def test_les_reperes_portent_les_numeros_du_plan(self):
        self.generate(champs=("1", "2", "3"))
        labels = [t.text for anchor in self.anchors() for t in anchor.iter(TEXT)]
        # Chaque forme est écrite deux fois : la moderne et son repli hérité.
        self.assertEqual(sorted(set(labels)), ["1", "2", "3"])

    def test_identifiants_de_formes_uniques(self):
        self.generate(champs=("1", "2", "3", "4", "5"))
        ids = [e.get("id") for e in self.document().element.body.iter(qn("wp:docPr"))]
        self.assertEqual(len(set(ids)), len(ids))

    def test_les_reperes_sont_alignes_en_rangee(self):
        self.generate(champs=("1", "2", "3"))
        horizontal = self.offsets()[::2]
        self.assertEqual(len(set(horizontal)), 3)

    def test_pas_de_reperes_sans_demande_du_plan(self):
        self.generate(block={"markers": None})
        self.assertEqual(self.anchors(), [])

    def test_pas_de_reperes_sans_element_a_numeroter(self):
        self.generate(champs=())
        self.assertEqual(self.anchors(), [])


class ARegenerationTest(Harness):
    """Ce que la régénération doit à ce qu'on a fait dans Word."""

    def move_first_marker(self, left: str = "1234567", top: str = "-987654") -> None:
        """Reproduit un repère glissé à la souris sur la capture."""
        document = self.document()
        anchor = next(document.element.body.iter(ANCHOR))
        horizontal, vertical = list(anchor.iter(OFFSET))[:2]
        horizontal.text, vertical.text = left, top
        document.save(self.path)

    def renumber_caption(self, number: str = "7") -> None:
        """Reproduit Word recalculant le champ après une capture supprimée."""
        document = self.document()
        paragraph = next(p for p in document.paragraphs if p.text.startswith("Figure"))
        run = next(t for t in paragraph._p.iter(TEXT) if t.text == "1")
        run.text = number
        document.save(self.path)

    def test_un_repere_deplace_reste_ou_on_l_a_mis(self):
        self.generate()
        self.move_first_marker()
        self.generate()
        self.assertIn("1234567", self.offsets())
        self.assertIn("-987654", self.offsets())

    def test_des_reperes_non_touches_suivent_le_plan(self):
        self.generate(champs=("1", "2"))
        self.generate(champs=("1", "2", "3"))
        self.assertEqual(len(self.anchors()), 3)

    def test_une_legende_renumerotee_par_word_n_est_pas_une_redaction(self):
        """
        Word a changé le numéro tout seul : la légende suit toujours le plan.

        Sans cette distinction, une capture supprimée figeait la légende de
        toutes les suivantes — leur description ne se mettait plus à jour.
        """
        self.generate()
        self.renumber_caption()
        self.generate(block={"description": "Capture du visuel refondu"})
        self.assertIn("Figure 1 — Capture du visuel refondu", self.texts())

    def test_une_legende_reecrite_a_la_main_est_conservee(self):
        self.generate()
        document = self.document()
        paragraph = next(p for p in document.paragraphs if p.text.startswith("Figure"))
        for run in list(paragraph.runs)[1:]:
            run._r.getparent().remove(run._r)
        paragraph.runs[0].text = "Ma légende à moi"
        document.save(self.path)

        self.generate(block={"description": "Capture du visuel refondu"})
        self.assertIn("Ma légende à moi", self.texts())


if __name__ == "__main__":
    unittest.main()
