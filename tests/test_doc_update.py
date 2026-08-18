"""Tests de la mise à jour d'un document déjà rédigé."""

import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.doc_config import DocConfig
from src.generators.doc_updater import (
    CHANGED,
    NEW,
    REMOVED,
    ItemChange,
    UpdateReport,
    _element_span,
    _note_bookmark,
    _renumber_bookmarks,
    compare,
)
from src.generators.docx_index import index_document
from src.generators.word_generator import block_bookmark

# L'index travaille sur les identifiants de style, pas sur leurs noms.
HEADINGS = {"Heading1": 1, "Heading2": 2, "Heading3": 3}


class FakeDoc:
    """Petit document construit à la main, indexable comme un vrai."""

    def __init__(self):
        self.document = Document()
        self.element = self.document.element
        self._bookmark_id = 0

    def heading(self, text: str, level: int, bookmark: str | None = None):
        paragraph = self.document.add_paragraph(text, style=f"Heading {level}")
        if bookmark:
            self._bookmark_id += 1
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), str(self._bookmark_id))
            start.set(qn("w:name"), bookmark)
            paragraph._p.insert(0, start)
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), str(self._bookmark_id))
            paragraph._p.append(end)
        return paragraph

    def block(self, bookmark: str, *lines: str, style: str = "Normal"):
        paragraphs = [self.document.add_paragraph(line, style=style) for line in lines]
        self._bookmark_id += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), bookmark)
        paragraphs[0]._p.addprevious(start)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        paragraphs[-1]._p.addnext(end)
        return paragraphs


def sample(expression: str, description: str, todo: bool = False):
    """Document minimal : une mesure, son code DAX et sa description."""
    doc = FakeDoc()
    doc.heading("Définition des mesures", 1)
    doc.heading("Chiffre d'affaires", 3, "measure_ca")
    doc.block("measure_ca_code", "Code DAX", expression)
    doc.block(
        "measure_ca_desc", "Description", description, style="List Bullet" if todo else "Normal"
    )
    return doc


BLOCKS = {"measure_ca_code": "mesure_code", "measure_ca_desc": "mesure_description"}


class IndexTest(unittest.TestCase):
    def test_item_et_blocs_reperes(self):
        index = index_document(sample("SUM(x)", "texte"), HEADINGS, BLOCKS)

        item = index.get("measure_ca")
        self.assertIsNotNone(item)
        self.assertEqual(item.title, "Chiffre d'affaires")
        self.assertEqual(sorted(item.blocks), ["mesure_code", "mesure_description"])
        self.assertIn("SUM(x)", item.blocks["mesure_code"].text)

    def test_bloc_attribue_au_titre_le_plus_profond(self):
        doc = FakeDoc()
        doc.heading("Ventes", 2, "groupe_ventes")
        doc.heading("Chiffre d'affaires", 3, "measure_ca")
        doc.block("measure_ca_code", "Code DAX", "SUM(x)")

        index = index_document(doc, HEADINGS, BLOCKS)
        self.assertIn("mesure_code", index.get("measure_ca").blocks)
        self.assertNotIn("mesure_code", index.get("groupe_ventes").blocks)

    def test_suffixe_technique_ignore_dans_le_titre(self):
        doc = FakeDoc()
        paragraph = doc.heading("CA du mois", 3, "visual_v1")
        run = paragraph.add_run("   card")
        run.style = doc.document.styles["Emphasis"]

        index = index_document(doc, HEADINGS, {}, {"Emphasis"})
        self.assertEqual(index.get("visual_v1").title, "CA du mois")


class CompareTest(unittest.TestCase):
    def _compare(self, before, after, review=frozenset(), todo_style=""):
        return compare(
            index_document(before, HEADINGS, BLOCKS),
            index_document(after, HEADINGS, BLOCKS),
            set(review),
            todo_style,
        )

    def test_aucun_changement(self):
        report = self._compare(sample("SUM(x)", "texte"), sample("SUM(x)", "texte"))
        self.assertTrue(report.empty)

    def test_formule_modifiee(self):
        report = self._compare(sample("SUM(x)", "texte"), sample("SUMX(x)", "texte"))
        self.assertEqual([c.status for c in report.changed], [CHANGED])
        self.assertEqual(report.changed[0].blocks, ["mesure_code"])

    def test_texte_utilisateur_ignore(self):
        report = self._compare(
            sample("SUM(x)", "Texte rédigé par l'utilisateur"),
            sample("SUM(x)", "[À compléter]"),
            review={"mesure_description"},
        )
        self.assertTrue(report.empty)

    def test_bloc_a_completer_non_compare(self):
        before = sample("SUM(x)", "Texte rédigé")
        after = sample("SUM(x)", "[À compléter]", todo=True)
        report = self._compare(before, after, todo_style="ListBullet")
        self.assertTrue(report.empty)

    def test_item_ajoute(self):
        after = sample("SUM(x)", "texte")
        after.heading("Panier moyen", 3, "measure_panier")
        report = self._compare(sample("SUM(x)", "texte"), after)
        self.assertEqual([c.title for c in report.added], ["Panier moyen"])
        self.assertEqual(report.added[0].status, NEW)

    def test_item_disparu(self):
        before = sample("SUM(x)", "texte")
        before.heading("Ancienne mesure", 3, "measure_vieille")
        report = self._compare(before, sample("SUM(x)", "texte"))
        self.assertEqual([c.title for c in report.removed], ["Ancienne mesure"])
        self.assertEqual(report.removed[0].status, REMOVED)


class ReportTest(unittest.TestCase):
    def test_libelle_lisible(self):
        change = ItemChange(
            "bm",
            "Chiffre d'affaires",
            CHANGED,
            ["mesure_code"],
            kind="mesure",
            block_labels={"mesure_code": "le code DAX"},
        )
        self.assertEqual(change.describe(), "mesure « Chiffre d'affaires » — modifié (le code DAX)")

    def test_document_a_jour(self):
        self.assertEqual(UpdateReport().lines(), ["  Aucun changement : le document est à jour"])

    def test_note_seule_compte_comme_modification(self):
        report = UpdateReport(notes=1)
        self.assertTrue(report.applied)
        self.assertTrue(report.empty)


class HelpersTest(unittest.TestCase):
    def test_span_inclut_les_marqueurs(self):
        doc = FakeDoc()
        doc.heading("Mesure", 3, "measure_ca")
        doc.block("measure_ca_code", "Code DAX", "SUM(x)")

        item = index_document(doc, HEADINGS, BLOCKS).get("measure_ca")
        span = _element_span(item.elements)
        tags = {node.tag for node in span}
        self.assertIn(qn("w:bookmarkStart"), tags)
        self.assertIn(qn("w:bookmarkEnd"), tags)

    def test_renumerotation_des_signets(self):
        doc = FakeDoc()
        doc.heading("A", 3, "item_a")
        doc.heading("B", 3, "item_b")
        for start in doc.element.body.iter(qn("w:bookmarkStart")):
            start.set(qn("w:id"), "1")  # collision volontaire
        for end in doc.element.body.iter(qn("w:bookmarkEnd")):
            end.set(qn("w:id"), "1")

        _renumber_bookmarks(doc.element.body)

        ids = [b.get(qn("w:id")) for b in doc.element.body.iter(qn("w:bookmarkStart"))]
        self.assertEqual(len(ids), len(set(ids)))

    def test_signet_de_note_stable(self):
        self.assertEqual(_note_bookmark("measure:CA"), _note_bookmark("measure:CA"))
        self.assertNotEqual(_note_bookmark("measure:CA"), _note_bookmark("measure:CA HT"))

    def test_nom_de_signet_de_bloc(self):
        self.assertEqual(block_bookmark("measure:CA", "mesure_code"), "measure:CA#mesure_code")


class ConfigTest(unittest.TestCase):
    def test_blocs_suivis_du_plan(self):
        config = DocConfig(
            {
                "sections": [
                    {
                        "id": "mesures",
                        "blocks": [
                            {
                                "type": "loop",
                                "section": {
                                    "blocks": [
                                        {"id": "code", "type": "property", "track": True},
                                        {"id": "desc", "type": "property", "review": True},
                                        {"id": "libre", "type": "paragraph"},
                                    ]
                                },
                            }
                        ],
                    }
                ]
            }
        )
        tracked, review = config.tracked_block_ids()
        self.assertEqual(tracked, {"code"})
        self.assertEqual(review, {"desc"})


if __name__ == "__main__":
    unittest.main()
