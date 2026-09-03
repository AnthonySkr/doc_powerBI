"""
Garde-fous du repérage : identifiants uniques, encadrements non imbriqués.

Deux pièges silencieux du plan YAML. Ils ne produisaient aucune erreur, mais
faisaient perdre la rédaction de l'utilisateur à la génération suivante.
"""

import unittest

from src import console
from src.config import DocConfig
from src.generators.word.merging import MergeWriter
from src.merge import markers


def writer(**merge_options):
    from docx import Document

    document = Document()
    config = DocConfig({"merge": {"enabled": True, **merge_options}})
    return document, MergeWriter(document, config, None)


def anchors(document) -> list[str]:
    return [
        marker.value
        for paragraph in document.paragraphs
        if (marker := markers.parse(paragraph.text)) is not None and marker.kind == markers.ELEMENT
    ]


def kinds(document) -> list[str]:
    return [
        marker.kind
        for paragraph in document.paragraphs
        if (marker := markers.parse(paragraph.text)) is not None
    ]


class IdentifiantsUniquesTest(unittest.TestCase):
    """Un `id:` de section placé dans une boucle se répétait à chaque tour."""

    def test_identifiants_distincts_inchanges(self):
        document, merge = writer()
        with console.silenced():
            merge.anchor({"id": "fiche"}, {})
            merge.anchor({"id": "autre"}, {})
        self.assertEqual(anchors(document), ["section:fiche", "section:autre"])

    def test_identifiant_repete_distingue_par_le_titre(self):
        """Le titre vient de la donnée parcourue : il ne bouge pas si elle est retriée."""
        document, merge = writer()
        with console.silenced():
            for name in ("Ventes", "Achats"):
                merge.anchor({"id": "fiche", "title": name}, {})
        self.assertEqual(anchors(document), ["section:fiche", "section:fiche>Achats"])

    def test_titres_identiques_restent_distincts(self):
        """
        Deux éléments parcourus peuvent porter le même titre.

        Le titre ne suffit alors plus à les distinguer : sans le rang en
        dernier ressort, deux ancres porteraient le même identifiant et la
        relecture ne saurait plus à laquelle rendre la rédaction.
        """
        document, merge = writer()
        with console.silenced():
            for _ in range(3):
                merge.anchor({"id": "fiche", "title": "Ventes"}, {})

        written = anchors(document)
        self.assertEqual(len(set(written)), len(written))

    def test_ordre_des_iterations_sans_effet_sur_le_nombre_d_ancres(self):
        """Quel que soit l'ordre, chaque élément parcouru reçoit son ancre."""
        first, second = writer(), writer()
        with console.silenced():
            for (_, merge), names in (
                (first, ("Ventes", "Achats")),
                (second, ("Achats", "Ventes")),
            ):
                for name in names:
                    merge.anchor({"id": "fiche", "title": name}, {})

        for document, _ in (first, second):
            written = anchors(document)
            self.assertEqual(len(written), 2)
            self.assertEqual(len(set(written)), 2)

    def test_identifiant_repete_sans_titre_numerote(self):
        """Sans titre il ne reste que le rang : un pis-aller, mais rien n'est perdu."""
        document, merge = writer()
        with console.silenced():
            for _ in range(3):
                merge.anchor({"id": "fiche"}, {})
        self.assertEqual(anchors(document), ["section:fiche", "section:fiche#2", "section:fiche#3"])

    def test_repetition_signalee_une_seule_fois(self):
        _document, merge = writer()
        messages: list[str] = []
        original = console.warn
        console.warn = messages.append
        try:
            for _ in range(4):
                merge.anchor({"id": "fiche"}, {})
        finally:
            console.warn = original
        self.assertEqual(len(messages), 1)
        self.assertIn("bookmark", messages[0])

    def test_bookmark_rendu_avant_deduplication(self):
        """Un bookmark bâti sur la donnée parcourue reste distinct : rien à suffixer."""
        document, merge = writer()
        with console.silenced():
            for name in ("Marge", "Chiffre d'affaires"):
                merge.anchor({"bookmark": "measure:{{ m.name }}"}, {"m": Named(name)})
        self.assertEqual(anchors(document), ["measure:Marge", "measure:Chiffre d'affaires"])


class Named:
    def __init__(self, name: str):
        self.name = name


class EncadrementsNonImbriquesTest(unittest.TestCase):
    """Un `gen` dans un `gen` laissait un `endgen` orphelin au découpage."""

    def test_bloc_simple_encadre(self):
        document, merge = writer()
        with console.silenced(), merge.delimit({"id": "table", "type": "table"}):
            document.add_paragraph("donnée")
        self.assertEqual(kinds(document), [markers.GENERATED, markers.GENERATED_END])

    def test_bloc_imbrique_non_encadre(self):
        document, merge = writer()
        with (  # noqa: SIM117
            console.silenced(),
            merge.delimit({"id": "dehors", "generated": True}),
        ):
            with merge.delimit({"id": "dedans", "type": "table"}):
                document.add_paragraph("donnée")
        self.assertEqual(kinds(document), [markers.GENERATED, markers.GENERATED_END])

    def test_bloc_imbrique_signale(self):
        document, merge = writer()
        messages: list[str] = []
        original = console.warn
        console.warn = messages.append
        try:
            with merge.delimit({"id": "dehors", "generated": True}):  # noqa: SIM117
                with merge.delimit({"id": "dedans", "type": "table"}):
                    document.add_paragraph("donnée")
        finally:
            console.warn = original
        self.assertEqual(len(messages), 1)
        self.assertIn("dedans", messages[0])

    def test_encadrement_reprend_apres_un_bloc_imbrique(self):
        """Le compteur redescend : le bloc suivant est de nouveau encadré."""
        document, merge = writer()
        with console.silenced():
            with merge.delimit({"id": "dehors", "generated": True}):  # noqa: SIM117
                with merge.delimit({"id": "dedans", "type": "table"}):
                    document.add_paragraph("donnée")
            with merge.delimit({"id": "apres", "type": "table"}):
                document.add_paragraph("autre")
        self.assertEqual(
            kinds(document),
            [
                markers.GENERATED,
                markers.GENERATED_END,
                markers.GENERATED,
                markers.GENERATED_END,
            ],
        )

    def test_marqueurs_exclus_des_empreintes(self):
        """Les empreintes ne portent que sur le contenu, jamais sur les marqueurs."""
        document, merge = writer()
        with console.silenced(), merge.delimit({"id": "table", "type": "table"}):
            document.add_paragraph("donnée")
        end = markers.parse(document.paragraphs[-1].text)
        assert end is not None
        self.assertEqual(end.digests, (markers.fingerprint("donnée"),))


if __name__ == "__main__":
    unittest.main()
