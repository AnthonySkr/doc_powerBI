"""Tests de la mise en page des tableaux et des styles du template v2."""

import unittest

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from src.generators.word.tables import (
    keep_row_together,
    repeat_header_row,
    set_fixed_layout,
    set_table_look,
    set_vertical_align,
)
from src.models.data_models import VisualReference


class FixedLayoutTest(unittest.TestCase):
    def _table(self, columns=3):
        document = Document()
        return document.add_table(rows=1, cols=columns)

    def test_grille_alignee_sur_les_largeurs(self):
        table = self._table()
        set_fixed_layout(table, [1.25, 3.5, 11.25])

        widths = [g.get(qn("w:w")) for g in table._tbl.iter(qn("w:gridCol"))]
        self.assertEqual(widths, [str(Cm(w).twips) for w in (1.25, 3.5, 11.25)])

    def test_layout_fixe(self):
        table = self._table()
        set_fixed_layout(table, [1.25, 3.5, 11.25])

        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        self.assertEqual(layout.get(qn("w:type")), "fixed")

    def test_colonne_sans_largeur_inchangee(self):
        table = self._table()
        before = table._tbl.find(qn("w:tblGrid"))[1].get(qn("w:w"))
        set_fixed_layout(table, [1.25, None, 11.25])
        self.assertEqual(table._tbl.find(qn("w:tblGrid"))[1].get(qn("w:w")), before)

    def test_appels_repetes_sans_doublon(self):
        table = self._table()
        set_fixed_layout(table, [1.25, 3.5, 11.25])
        set_fixed_layout(table, [1.25, 3.5, 11.25])
        self.assertEqual(len(table._tbl.tblPr.findall(qn("w:tblLayout"))), 1)


class TableLookTest(unittest.TestCase):
    def test_ligne_d_entete_mise_en_valeur(self):
        table = Document().add_table(rows=1, cols=2)
        set_table_look(table, first_row=True)

        look = table._tbl.tblPr.find(qn("w:tblLook"))
        self.assertEqual(look.get(qn("w:firstRow")), "1")
        self.assertEqual(look.get(qn("w:noVBand")), "1")

    def test_sans_entete(self):
        table = Document().add_table(rows=1, cols=2)
        set_table_look(table, first_row=False)
        self.assertEqual(table._tbl.tblPr.find(qn("w:tblLook")).get(qn("w:firstRow")), "0")


class RowPropertiesTest(unittest.TestCase):
    def test_entete_repete(self):
        row = Document().add_table(rows=1, cols=1).rows[0]
        repeat_header_row(row)
        repeat_header_row(row)
        self.assertEqual(len(row._tr.trPr.findall(qn("w:tblHeader"))), 1)

    def test_ligne_non_coupee(self):
        row = Document().add_table(rows=1, cols=1).rows[0]
        keep_row_together(row)
        self.assertIsNotNone(row._tr.trPr.find(qn("w:cantSplit")))

    def test_alignement_vertical(self):
        cell = Document().add_table(rows=1, cols=1).rows[0].cells[0]
        set_vertical_align(cell, "center")
        set_vertical_align(cell, "center")
        aligns = cell._tc.tcPr.findall(qn("w:vAlign"))
        self.assertEqual([a.get(qn("w:val")) for a in aligns], ["center"])


class VisualReferenceValueTest(unittest.TestCase):
    def test_element_expose_son_nom(self):
        ref = VisualReference(number="1", kind="mesure", name="CA", label="…", role="Valeur")
        self.assertEqual(ref.value, "CA")

    def test_filtre_expose_son_expression(self):
        ref = VisualReference(
            number="2",
            kind="filtre",
            name="Ventes.Actif",
            label="…",
            role="Filtre",
            expression="Ventes.Actif (Inclut: Oui)",
        )
        self.assertEqual(ref.value, "Ventes.Actif (Inclut: Oui)")


if __name__ == "__main__":
    unittest.main()
