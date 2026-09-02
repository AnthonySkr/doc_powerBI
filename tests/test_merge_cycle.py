"""
Cycle complet : génération, remaniement libre par l'utilisateur, regénération.

Le contrat vérifié ici : le script réécrit ses données, et ne touche à rien
d'autre — titre reformulé, note ajoutée, capture collée, paragraphes en plus,
y compris ce qui a été écrit *à l'intérieur* d'un contenu du script (sous son
tableau, entre ses valeurs).
"""

import os
import re
import struct
import tempfile
import unittest
import zlib
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph

from src import console
from src.config import DocConfig
from src.generators.word import generate_word_documentation
from src.merge import markers, orphans
from src.models.data_models import DaxMeasure, MeasureGroup, SemanticModel

_DRAWING = qn("w:drawing")

# Le plan livré ne surligne plus rien ; les tests du mécanisme le redemandent.
_HIGHLIGHTED = {"merge": {"highlight_changed": "yellow", "highlight_new": "green"}}

PLAN = {
    "document": {"template": "", "cover": {}, "properties": {}},
    "rendering": {"page_break_before_heading_1": False, "links": {"enabled": False}},
    "sections": [
        {
            "id": "mesures",
            "title": "Mesures",
            "level": 1,
            "blocks": [
                {
                    "type": "loop",
                    "over": "model.tables_with_measures",
                    "item": "groupe",
                    "section": {
                        "title": "{{ groupe.name }}",
                        "level": 2,
                        "blocks": [
                            {
                                "type": "loop",
                                "over": "groupe.measures",
                                "item": "measure",
                                "section": {
                                    "title": "{{ measure.name }}",
                                    "level": 3,
                                    "bookmark": "measure:{{ measure.name }}",
                                    "fingerprint": "{{ measure.expression }}",
                                    "blocks": [
                                        {
                                            "type": "property",
                                            "id": "code",
                                            "label": "Code DAX",
                                            "value": "{{ measure.expression }}",
                                        },
                                        {
                                            "type": "property",
                                            "id": "sources",
                                            "label": "Sources",
                                            "value_list": "{{ measure.used_columns }}",
                                        },
                                        {
                                            "type": "table",
                                            "id": "champs",
                                            "label": "Champs",
                                            "over": "groupe.measures",
                                            "item": "champ",
                                            "header": True,
                                            "header_labels": ["Mesure"],
                                            "columns": [{"id": "nom", "value": "{{ champ.name }}"}],
                                        },
                                        {"type": "user_fill", "id": "commentaire"},
                                    ],
                                },
                            }
                        ],
                    },
                }
            ],
        }
    ],
}


def png(path: str) -> str:
    """Écrit une image PNG minimale mais valide."""

    def chunk(kind: bytes, data: bytes) -> bytes:
        body = kind + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body))

    raw = b"".join(b"\x00" + bytes([200, 30, 60] * 4) for _ in range(4))
    with open(path, "wb") as f:
        f.write(
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", 4, 4, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", zlib.compress(raw))
            + chunk(b"IEND", b"")
        )
    return path


def measure_section(plan: dict) -> dict:
    """La section d'une mesure dans une copie du plan de référence."""
    return plan["sections"][0]["blocks"][0]["section"]["blocks"][0]["section"]


def context(**expressions: str) -> dict:
    measures = [
        DaxMeasure(
            name=name,
            expression=expression,
            table_name="Ventes",
            # Colonnes citées par l'expression : une liste de valeurs qui bouge
            # d'une génération à l'autre, comme dans le plan livré.
            used_columns=sorted(set(re.findall(r"\w+\[[^\]]+\]", expression))),
        )
        for name, expression in expressions.items()
    ]
    return {
        "model": SemanticModel(
            tables_with_measures=[MeasureGroup(name="Ventes", measures=measures)]
        ),
        "report": None,
        "inputs": {},
        "styles": {},
    }


class MergeHarness(unittest.TestCase):
    """
    Génère, laisse remanier le document comme le ferait un utilisateur, et
    régénère. Les gestes sont ceux qu'on fait dans Word : reformuler un titre,
    ajouter un paragraphe, coller une capture, écrire sous un tableau.
    """

    def setUp(self):
        self._directory = tempfile.TemporaryDirectory()
        self.addCleanup(self._directory.cleanup)
        self.path = os.path.join(self._directory.name, "doc.docx")
        self.template = os.path.join(self._directory.name, "template.docx")
        Document().save(self.template)

    # ── Utilitaires ───────────────────────────────────────────────
    def generate(self, overrides: dict | None = None, **expressions: str):
        raw = {**PLAN, **(overrides or {})}
        raw["document"] = {**raw["document"], "template": self.template}
        with console.silenced():
            return generate_word_documentation(DocConfig(raw), context(**expressions), self.path)

    def texts(self) -> list[str]:
        return [p.text for p in Document(self.path).paragraphs if p.text.strip()]

    def _annexe_start(self, document) -> int:
        """Rang, dans le corps, de l'ancre de l'annexe — la fin du corps sinon."""
        body = list(document.element.body)
        for rank, node in enumerate(body):
            marker = markers.of(node)
            if marker is not None and marker.value == orphans.ELEMENT_ID:
                return rank
        return len(body)

    def _split_annexe(self) -> tuple[list[str], list[str]]:
        """Le document, de part et d'autre de l'annexe des contenus non replacés."""
        texts = self.texts()
        marker = f"pbi::elem|{orphans.ELEMENT_ID}|"
        start = next((i for i, text in enumerate(texts) if text.startswith(marker)), len(texts))
        return texts[:start], texts[start:]

    def before_annexe(self) -> list[str]:
        return self._split_annexe()[0]

    def annexe(self) -> list[str]:
        return self._split_annexe()[1]

    def annexe_content(self) -> list[str]:
        """Tout ce que porte l'annexe, cellules de tableau comprises."""
        document = Document(self.path)
        body = list(document.element.body)
        return [node.xpath("string(.)") for node in body[self._annexe_start(document) :]]

    def drop_annexe(self) -> None:
        """Supprime l'annexe, comme le ferait l'utilisateur une fois reclassée."""
        document = Document(self.path)
        body = document.element.body
        for node in list(body)[self._annexe_start(document) :]:
            body.remove(node)
        document.save(self.path)

    def clear_runs(self, paragraph) -> None:
        """Vide un paragraphe de son texte, sans toucher à ce qui l'entoure."""
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)

    def find(self, needle: str):
        document = Document(self.path)
        return document, next(p for p in document.paragraphs if needle in p.text)

    def rewrite(self, needle: str, text: str) -> None:
        """Remplace le texte d'un paragraphe, comme le ferait l'utilisateur."""
        document, paragraph = self.find(needle)
        self.clear_runs(paragraph)
        paragraph.add_run(text)
        document.save(self.path)

    def add_after(self, needle: str, text: str) -> None:
        """Insère un paragraphe supplémentaire après un autre."""
        document, paragraph = self.find(needle)
        node = deepcopy(paragraph._p)
        paragraph._p.addnext(node)
        added = Paragraph(node, paragraph._parent)
        self.clear_runs(added)
        added.add_run(text)
        document.save(self.path)

    def write_under_table(self, text: str) -> None:
        """
        Écrit sous le tableau, dans la ligne que le script laisse là.

        C'est le geste naturel : on clique sous le tableau et on décrit ce
        qu'il montre.
        """
        document = Document(self.path)
        table = document.tables[0]
        paragraph = Paragraph(table._tbl.getnext(), document)
        paragraph.add_run(text)
        document.save(self.path)

    def add_above_table(self, text: str) -> None:
        """Insère un paragraphe entre le sous-titre du tableau et le tableau."""
        document = Document(self.path)
        table = document.tables[0]
        added = deepcopy(table._tbl.getnext())
        table._tbl.addprevious(added)
        Paragraph(added, document).add_run(text)
        document.save(self.path)

    def forget_digests(self) -> None:
        """
        Ramène les marqueurs de fin à leur forme d'avant les empreintes, pour
        rejouer un document produit par une version antérieure du script.
        """
        document = Document(self.path)
        for paragraph in document.paragraphs:
            if paragraph.text.startswith("pbi::endgen"):
                paragraph.runs[0].text = "pbi::endgen"
        document.save(self.path)

    def layout(self) -> list[str]:
        """Corps du document : textes visibles et tableaux, dans l'ordre."""
        document = Document(self.path)
        items = []
        for node in document.element.body:
            if node.tag == qn("w:tbl"):
                items.append("[tableau]")
                continue
            text = Paragraph(node, document).text if node.tag == qn("w:p") else ""
            if text.strip() and not text.startswith("pbi::"):
                items.append(text)
        return items

    def paste_image(self, needle: str) -> None:
        document, paragraph = self.find(needle)
        self.clear_runs(paragraph)
        image = png(os.path.join(self._directory.name, "capture.png"))
        paragraph.add_run().add_picture(image, width=Cm(2))
        document.save(self.path)

    def images(self) -> int:
        document = Document(self.path)
        drawings = document.element.body.findall(f".//{_DRAWING}")
        return sum(
            1
            for drawing in drawings
            for blip in drawing.iter(qn("a:blip"))
            if document.part.related_parts.get(blip.get(qn("r:embed"))) is not None
        )

    def highlight(self, needle: str):
        _, paragraph = self.find(needle)
        colors = {run.font.highlight_color for run in paragraph.runs if run.text.strip()} - {None}
        return next(iter(colors), None)


class MergeCycleTest(MergeHarness):
    """Le contrat : le script réécrit ses données, et ne touche à rien d'autre."""

    def test_premiere_generation(self):
        log = self.generate(CA="SUM(Ventes[Montant])")
        self.assertFalse(log.is_update)
        self.assertIn("measure:CA", log.written_ids)

    def test_titre_reformule_conserve(self):
        self.generate(CA="1")
        self.rewrite("CA", "CA — Chiffre d'affaires net")
        self.generate(CA="1")
        self.assertIn("CA — Chiffre d'affaires net", self.texts())

    def test_paragraphe_ajoute_conserve(self):
        self.generate(CA="1")
        self.add_after("[À compléter]", "Note ajoutée à la main.")
        self.generate(CA="1")
        self.assertIn("Note ajoutée à la main.", self.texts())

    def test_zone_a_completer_remplie_conservee(self):
        self.generate(CA="1")
        self.rewrite("[À compléter]", "Somme facturée sur la période.")
        self.generate(CA="1")
        self.assertIn("Somme facturée sur la période.", self.texts())
        self.assertNotIn("[À compléter]", self.texts())

    def test_capture_collee_conservee_avec_son_image(self):
        self.generate(CA="1")
        self.paste_image("[À compléter]")
        self.assertEqual(self.images(), 1)

        self.generate(CA="1")
        self.assertEqual(self.images(), 1)

    def test_donnees_du_script_mises_a_jour(self):
        self.generate(CA="SUM(Ventes[Montant])")
        self.rewrite("[À compléter]", "Explication.")

        self.generate(CA="SUM(Ventes[MontantHT])")

        self.assertIn("SUM(Ventes[MontantHT])", self.texts())
        self.assertNotIn("SUM(Ventes[Montant])", self.texts())
        self.assertIn("Explication.", self.texts())

    def test_document_sans_marque_quand_la_technique_change(self):
        """Par défaut le document ne signale rien : le résumé console s'en charge."""
        self.generate(CA="1")
        self.rewrite("[À compléter]", "Explication.")

        log = self.generate(CA="2")

        self.assertIsNone(self.highlight("Explication."))
        self.assertEqual([log.title_of(name) for name in log.changed], ["CA"])
        self.assertIn("· CA", log.details())

    def test_texte_signale_quand_la_technique_change(self):
        """Le surlignage reste disponible, sur demande explicite."""
        self.generate(CA="1")
        self.rewrite("[À compléter]", "Explication.")

        self.generate(_HIGHLIGHTED, CA="2")
        self.assertIsNotNone(self.highlight("Explication."))

    def test_signalement_retire_a_la_generation_suivante(self):
        self.generate(_HIGHLIGHTED, CA="1")
        self.rewrite("[À compléter]", "Explication.")
        self.generate(_HIGHLIGHTED, CA="2")
        self.generate(_HIGHLIGHTED, CA="2")
        self.assertIsNone(self.highlight("Explication."))

    # ── Écrire à l'intérieur d'un contenu du script ───────────────
    def test_description_sous_un_tableau_conservee(self):
        self.generate(CA="1")
        self.write_under_table("Ce tableau liste les mesures de la table.")

        self.generate(CA="1")

        self.assertIn("Ce tableau liste les mesures de la table.", self.texts())

    def test_description_sous_un_tableau_remise_a_sa_place(self):
        self.generate(CA="1")
        self.write_under_table("Lecture du tableau.")

        self.generate(CA="1")

        layout = self.layout()
        self.assertEqual(layout[layout.index("[tableau]") + 1], "Lecture du tableau.")

    def test_description_sous_un_tableau_conservee_a_chaque_generation(self):
        self.generate(CA="1")
        self.write_under_table("Lecture du tableau.")
        self.generate(CA="1")
        reference = self.layout()

        for _ in range(3):
            self.generate(CA="1")
        self.assertEqual(self.layout(), reference)

    def test_description_conservee_quand_le_tableau_change(self):
        self.generate(CA="1")
        self.write_under_table("Lecture du tableau.")

        self.generate(CA="1", Marge="2")

        self.assertIn("Lecture du tableau.", self.texts())
        self.assertEqual(self.texts().count("Lecture du tableau."), 1)

    def test_note_glissee_avant_un_tableau_conservee_a_sa_place(self):
        self.generate(CA="1")
        self.add_above_table("Avertissement avant le tableau.")

        self.generate(CA="1")

        layout = self.layout()
        self.assertEqual(layout[layout.index("[tableau]") - 1], "Avertissement avant le tableau.")

    def test_note_glissee_sous_le_code_dax_conservee(self):
        self.generate(CA="SUM(Ventes[Montant])")
        self.add_after("SUM(Ventes[Montant])", "Formule reprise du cahier des charges.")

        self.generate(CA="SUM(Ventes[Montant])")

        self.assertIn("Formule reprise du cahier des charges.", self.texts())

    def test_capture_collee_sous_un_tableau_conservee(self):
        self.generate(CA="1")
        document = Document(self.path)
        paragraph = Paragraph(document.tables[0]._tbl.getnext(), document)
        paragraph.add_run().add_picture(png(os.path.join(self._directory.name, "vue.png")), Cm(2))
        document.save(self.path)

        self.generate(CA="1")

        self.assertEqual(self.images(), 1)

    def test_donnee_du_script_reecrite_et_jamais_doublee(self):
        self.generate(CA="SUM(Ventes[Montant])")
        self.write_under_table("Lecture du tableau.")

        self.generate(CA="SUM(Ventes[MontantHT])")

        self.assertIn("SUM(Ventes[MontantHT])", self.texts())
        self.assertNotIn("SUM(Ventes[Montant])", self.texts())
        self.assertEqual(self.layout().count("[tableau]"), 1)

    def test_donnee_du_script_remaniee_a_la_main_est_reecrite(self):
        """Le script reste propriétaire de ses données : la sienne l'emporte."""
        self.generate(CA="SUM(Ventes[Montant])")
        self.rewrite("SUM(Ventes[Montant])", "SUM(Ventes[Autre])")

        self.generate(CA="SUM(Ventes[Montant])")

        self.assertIn("SUM(Ventes[Montant])", self.texts())
        self.assertNotIn("SUM(Ventes[Autre])", self.before_annexe())

    def test_donnee_remaniee_a_la_main_recueillie_en_annexe(self):
        """Réécrite à sa place, mais pas jetée : la version retouchée est gardée."""
        self.generate(CA="SUM(Ventes[Montant])")
        self.rewrite("SUM(Ventes[Montant])", "SUM(Ventes[Autre])")

        self.generate(CA="SUM(Ventes[Montant])")

        self.assertIn("SUM(Ventes[Autre])", self.annexe())

    def test_valeur_disparue_du_script_ne_revient_pas(self):
        self.generate(CA="SUM(Ventes[A]) + SUM(Ventes[B])")
        self.write_under_table("Lecture du tableau.")

        self.generate(CA="SUM(Ventes[A])")

        self.assertNotIn("Ventes[B]", self.texts())
        self.assertIn("Lecture du tableau.", self.texts())

    def test_document_d_une_version_anterieure_rend_ce_qui_suit_son_tableau(self):
        """Sans empreintes, on récupère au moins ce qui suit le tableau."""
        self.generate(CA="1")
        self.write_under_table("Écrit avec l'ancienne version.")
        self.forget_digests()

        self.generate(CA="1")

        self.assertIn("Écrit avec l'ancienne version.", self.texts())

    def test_titre_jamais_surligne(self):
        self.generate(CA="1")
        self.generate(CA="2")
        self.assertIsNone(self.highlight("CA"))

    def test_document_stable_sur_plusieurs_generations(self):
        self.generate(CA="1")
        self.rewrite("[À compléter]", "Explication.")
        self.generate(CA="1")
        reference = self.texts()

        for _ in range(3):
            self.generate(CA="1")
        self.assertEqual(self.texts(), reference)

    def test_nouvel_element_detecte(self):
        self.generate(CA="1")
        log = self.generate(CA="1", Marge="2")
        self.assertEqual(log.new, ["measure:Marge"])
        self.assertEqual(log.changed, [])

    def test_element_retire_signale(self):
        self.generate(CA="1", Marge="2")
        log = self.generate(CA="1")
        self.assertEqual(log.removed, ["measure:Marge"])

    def test_version_precedente_archivee(self):
        self.generate(CA="1")
        self.generate(CA="2")
        versions = os.listdir(os.path.join(self._directory.name, ".versions"))
        self.assertEqual(len(versions), 1)

    def test_fusion_desactivable(self):
        self.generate(CA="1")
        self.rewrite("[À compléter]", "Explication.")

        log = self.generate({"merge": {"enabled": False}}, CA="1")

        self.assertFalse(log.is_update)
        self.assertNotIn("Explication.", self.texts())
        self.assertIn("[À compléter]", self.texts())


if __name__ == "__main__":
    unittest.main()
