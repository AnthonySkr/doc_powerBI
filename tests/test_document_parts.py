"""Tests des en-têtes, de la table des matières et des liens retour."""

import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.doc_config import DocConfig
from src.generators.data_context import _index_usages
from src.generators.word_generator import (
    _mark_toc_fields_dirty,
    _replace_in_paragraph,
    _set_toc_levels,
    _set_update_fields,
)
from src.models.data_models import (
    DaxMeasure,
    PowerBIReport,
    ReportPage,
    Visual,
    VisualElement,
)


def toc_paragraph(document, instruction=' TOC \\o "1-2" \\h \\z \\u '):
    """Reproduit un champ TOC : fldChar begin / instrText / fldChar end."""
    paragraph = document.add_paragraph()
    for kind, text in (("begin", None), (None, instruction), ("end", None)):
        run = OxmlElement("w:r")
        if kind:
            char = OxmlElement("w:fldChar")
            char.set(qn("w:fldCharType"), kind)
            run.append(char)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run.append(instr)
        paragraph._p.append(run)
    return paragraph


class ReplaceInParagraphTest(unittest.TestCase):
    def test_remplacement_dans_un_seul_run(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("\tTitre intercalaire 1")

        self.assertEqual(_replace_in_paragraph(paragraph, "Titre intercalaire 1", "Ventes"), 1)
        self.assertEqual(paragraph.text, "\tVentes")

    def test_tabulation_conservee(self):
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run._r.append(OxmlElement("w:tab"))
        run.add_text("Titre intercalaire 1")

        _replace_in_paragraph(paragraph, "Titre intercalaire 1", "Ventes")
        self.assertEqual(len(paragraph._p.findall(".//" + qn("w:tab"))), 1)
        self.assertIn("Ventes", paragraph.text)

    def test_texte_reparti_sur_plusieurs_runs(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Titre inter")
        paragraph.add_run("calaire 1")

        self.assertEqual(_replace_in_paragraph(paragraph, "Titre intercalaire 1", "Ventes"), 1)
        self.assertEqual(paragraph.text, "Ventes")

    def test_texte_absent(self):
        document = Document()
        paragraph = document.add_paragraph("Autre chose")
        self.assertEqual(_replace_in_paragraph(paragraph, "Titre", "Ventes"), 0)
        self.assertEqual(paragraph.text, "Autre chose")


class TableOfContentsTest(unittest.TestCase):
    def test_champ_marque_a_recalculer(self):
        document = Document()
        toc_paragraph(document)

        self.assertEqual(_mark_toc_fields_dirty(document.element.body), 1)
        dirty = [
            fc.get(qn("w:dirty"))
            for fc in document.element.body.iter(qn("w:fldChar"))
            if fc.get(qn("w:fldCharType")) == "begin"
        ]
        self.assertEqual(dirty, ["true"])

    def test_niveaux_modifies(self):
        document = Document()
        toc_paragraph(document)

        _mark_toc_fields_dirty(document.element.body, "1-3")
        instructions = [i.text for i in document.element.body.iter(qn("w:instrText"))]
        self.assertIn('\\o "1-3"', instructions[0])

    def test_pageref_ignore(self):
        document = Document()
        toc_paragraph(document, " PAGEREF _Toc236231224 \\h ")
        self.assertEqual(_mark_toc_fields_dirty(document.element.body), 0)

    def test_niveaux_sans_option(self):
        self.assertEqual(_set_toc_levels(' TOC \\o "1-2" \\h ', "1-3"), ' TOC \\o "1-3" \\h ')

    def test_update_fields_respecte_l_ordre_du_schema(self):
        document = Document()
        settings = document.settings.element
        _set_update_fields(settings)

        tags = [child.tag.split("}")[1] for child in settings]
        self.assertIn("updateFields", tags)
        if "compat" in tags:
            self.assertLess(tags.index("updateFields"), tags.index("compat"))

    def test_update_fields_idempotent(self):
        document = Document()
        settings = document.settings.element
        _set_update_fields(settings)
        _set_update_fields(settings)
        self.assertEqual(len(settings.findall(qn("w:updateFields"))), 1)


class UsagesTest(unittest.TestCase):
    def _report(self):
        element = VisualElement(
            query_ref="Ventes.CA",
            display_name="CA total",
            friendly_name="CA total",
            type_category="Mesure",
            role="Values",
            property_name="CA",
        )
        visual = Visual(id="v1", visual_type="card", title="Carte CA", elements=[element])
        page = ReportPage(name="p1", display_name="Synthèse", visuals=[visual])
        return PowerBIReport(name="Demo", pages=[page])

    def test_usages_pointent_vers_le_visuel(self):
        measures = {"CA": DaxMeasure(name="CA", expression="1", table_name="Ventes")}
        options = DocConfig({}).section_options("visuels")

        _index_usages(self._report(), measures, options)

        usages = measures["CA"].usages
        self.assertEqual([u.text for u in usages], ["Synthèse — Carte CA"])
        self.assertEqual([u.target for u in usages], ["visual:p1:v1"])

    def test_mesure_inconnue_ignoree(self):
        measures: dict[str, DaxMeasure] = {}
        _index_usages(self._report(), measures, {})
        self.assertEqual(measures, {})

    def test_gabarits_personnalises(self):
        measures = {"CA": DaxMeasure(name="CA", expression="1", table_name="Ventes")}
        options = {
            "usages": {
                "target": "v:{{ visual.id }}",
                "label": "{{ visual.title }} ({{ page.display_name }})",
            }
        }

        _index_usages(self._report(), measures, options)

        usage = measures["CA"].usages[0]
        self.assertEqual(usage.text, "Carte CA (Synthèse)")
        self.assertEqual(usage.target, "v:v1")


if __name__ == "__main__":
    unittest.main()
