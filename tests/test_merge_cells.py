"""
Annoter un tableau du script de l'intérieur.

Le tableau est à lui — il le réécrit — mais ce qu'on ajoute **sous** sa donnée,
dans la même cellule, est à nous et doit retrouver sa place. Le rattachement se
fait par les données que porte la ligne : si elles ont bougé, l'annotation ne
commente plus la même chose et part en annexe plutôt que d'être reposée au
hasard.
"""

import unittest

from docx import Document

from tests.test_merge_cycle import MergeHarness


class CellAnnotationTest(MergeHarness):
    def annotate(self, row: int, column: int, text: str) -> None:
        """Écrit sous la donnée d'une cellule, comme on le fait dans Word."""
        document = Document(self.path)
        cell = document.tables[0].rows[row].cells[column]
        cell.add_paragraph(text)
        document.save(self.path)

    def annexe_texts(self) -> list[str]:
        """Tout ce que porte l'annexe, cellules de tableau comprises."""
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        from src.merge import orphans

        document = Document(self.path)
        found: list[str] = []
        inside = False
        for node in document.element.body:
            if node.tag == qn("w:p"):
                text = Paragraph(node, document).text
                if text.startswith(f"pbi::elem|{orphans.ELEMENT_ID}|"):
                    inside = True
            if inside:
                found.append(node.xpath("string(.)"))
        return found

    def cell_texts(self) -> list[str]:
        document = Document(self.path)
        return [
            paragraph.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        ]

    # ── Ce qui retrouve sa cellule ────────────────────────────────
    def test_annotation_conservee_dans_sa_cellule(self):
        self.generate(marge="SUM(T[a])")
        self.annotate(1, 0, "Attention : hors taxes")

        self.generate(marge="SUM(T[a])")

        self.assertIn("Attention : hors taxes", self.cell_texts())

    def test_annotation_non_recueillie_en_annexe(self):
        self.generate(marge="SUM(T[a])")
        self.annotate(1, 0, "Attention : hors taxes")

        self.generate(marge="SUM(T[a])")

        self.assertEqual(self.annexe(), [])

    def test_donnee_du_script_toujours_a_jour(self):
        self.generate(marge="SUM(T[a])")
        self.annotate(1, 0, "Ma note")

        self.generate(marge="SUM(T[b])")

        texts = self.cell_texts()
        self.assertIn("Ma note", texts)
        self.assertIn("marge", texts)

    def test_annotation_stable_sur_plusieurs_generations(self):
        self.generate(marge="SUM(T[a])")
        self.annotate(1, 0, "Ma note")
        self.generate(marge="SUM(T[a])")
        self.generate(marge="SUM(T[a])")

        self.assertEqual(self.cell_texts().count("Ma note"), 1)

    def test_plusieurs_annotations(self):
        self.generate(marge="SUM(T[a])", ca="SUM(T[b])")
        self.annotate(1, 0, "Première note")
        self.annotate(2, 0, "Seconde note")

        self.generate(marge="SUM(T[a])", ca="SUM(T[b])")

        texts = self.cell_texts()
        self.assertIn("Première note", texts)
        self.assertIn("Seconde note", texts)

    # ── Ce qui ne peut pas être rattaché ──────────────────────────
    def test_annotation_d_une_ligne_disparue_recueillie(self):
        """Sa ligne n'existe plus : la reposer ailleurs serait un contresens."""
        self.generate(marge="SUM(T[a])", ca="SUM(T[b])")
        # Le tableau de « ca » liste les deux mesures : on annote la ligne
        # « marge », qui disparaîtra du tableau au tour suivant.
        document = Document(self.path)
        row = next(
            row for row in document.tables[1].rows if row.cells[0].paragraphs[0].text == "marge"
        )
        row.cells[0].add_paragraph("Note sur cette ligne")
        document.save(self.path)

        self.generate(ca="SUM(T[b])")

        self.assertTrue(any("Note sur cette ligne" in text for text in self.annexe_texts()))

    def test_retouche_de_la_donnee_elle_meme_recueillie(self):
        """Écrire dans la ligne du script reste une réécriture de sa donnée."""
        self.generate(marge="SUM(T[a])")
        document = Document(self.path)
        cell = document.tables[0].rows[1].cells[0]
        cell.paragraphs[0].add_run(" — à vérifier")
        document.save(self.path)

        self.generate(marge="SUM(T[a])")

        self.assertTrue(any("à vérifier" in text for text in self.annexe_texts()))


if __name__ == "__main__":
    unittest.main()
