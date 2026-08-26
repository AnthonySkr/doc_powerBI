"""
Amorces des zones à rédiger : l'indication du plan plutôt que « [À compléter] ».

Le contrat vérifié ici : un bloc `user_fill` qui dit ce qu'on attend de lui
(`hint:`) l'écrit à la place de l'amorce générique, dans la même mise en forme
— et l'indication rejoint les documents déjà générés tant que personne n'a
écrit dans la zone.
"""

import copy
import os
import tempfile
import unittest

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from src import console
from src.config import DocConfig
from src.generators.word import generate_word_documentation

TODO_STYLE = "A completer"

PLAN = {
    "document": {"template": "", "cover": {}, "properties": {}},
    "styles": {"todo": TODO_STYLE, "normal": "Normal", "fallback": "Normal"},
    "rendering": {"page_break_before_heading_1": False, "links": {"enabled": False}},
    "sections": [
        {
            "id": "acquisition",
            "title": "Acquisition des données",
            "level": 1,
            "blocks": [{"type": "user_fill", "id": "contenu"}],
        }
    ],
}


class Harness(unittest.TestCase):
    """Génère une zone à rédiger, seule dans son document."""

    def setUp(self):
        self._directory = tempfile.TemporaryDirectory()
        self.addCleanup(self._directory.cleanup)
        self.path = os.path.join(self._directory.name, "doc.docx")
        self.template = os.path.join(self._directory.name, "template.docx")
        template = Document()
        template.styles.add_style(TODO_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        template.save(self.template)

    def generate(self, block: dict | None = None, rendering: dict | None = None) -> None:
        plan = copy.deepcopy(PLAN)
        plan["document"]["template"] = self.template
        if block:
            plan["sections"][0]["blocks"][0].update(block)
        if rendering:
            plan["rendering"]["user_fill"] = rendering
        context = {"model": None, "report": None, "inputs": {}, "styles": plan["styles"]}
        with console.silenced():
            generate_word_documentation(DocConfig(plan), context, self.path)

    def paragraphs(self) -> list:
        """Les paragraphes visibles : sans les marqueurs de fusion."""
        return [
            p
            for p in Document(self.path).paragraphs
            if p.text.strip() and not p.text.startswith("pbi::")
        ]

    def texts(self) -> list[str]:
        return [p.text for p in self.paragraphs()]

    def paragraph(self, text: str):
        return next(p for p in self.paragraphs() if p.text == text)


class UserFillHintTest(Harness):
    def test_sans_hint_amorce_generique(self):
        self.generate()
        self.assertIn("[À compléter]", self.texts())

    def test_hint_ecrit_a_la_place_de_l_amorce(self):
        self.generate({"hint": "Décrire l'origine des données"})
        self.assertIn("[Décrire l'origine des données]", self.texts())
        self.assertNotIn("[À compléter]", self.texts())

    def test_hint_dans_la_meme_mise_en_forme(self):
        self.generate()
        generic = self.paragraph("[À compléter]").style.name
        self.generate({"hint": "Décrire l'origine des données"})
        self.assertEqual(self.paragraph("[Décrire l'origine des données]").style.name, generic)
        self.assertEqual(generic, TODO_STYLE)

    def test_hint_suit_les_variables_du_plan(self):
        self.generate({"hint": "Décrire la source de {{ inputs.rapport }}"})
        self.assertIn("[Décrire la source de]", self.texts())

    def test_forme_de_l_indication_configurable(self):
        self.generate(
            {"hint": "Décrire l'origine des données"},
            {"hint_format": "→ {hint}", "placeholder_text": "[À compléter]"},
        )
        self.assertIn("→ Décrire l'origine des données", self.texts())

    def test_zone_volontairement_vide_le_reste(self):
        """`show_placeholder: false` demande une ligne vide : le hint n'y change rien."""
        self.generate({"hint": "Décrire l'origine des données", "show_placeholder": False})
        self.assertNotIn("[Décrire l'origine des données]", self.texts())

    def test_label_conserve_au_dessus_de_l_indication(self):
        self.generate({"label": "Lecture du visuel", "hint": "Ce que montre le visuel"})
        self.assertEqual(self.texts()[-2:], ["Lecture du visuel", "[Ce que montre le visuel]"])


class HintDansUnDocumentExistantTest(Harness):
    """L'amorce suit le plan tant que personne n'a écrit dedans."""

    def rewrite(self, text: str, replacement: str) -> None:
        document = Document(self.path)
        paragraph = next(p for p in document.paragraphs if p.text == text)
        for run in list(paragraph.runs)[1:]:
            run._r.getparent().remove(run._r)
        paragraph.runs[0].text = replacement
        document.save(self.path)

    def test_hint_ajoute_atteint_un_document_deja_genere(self):
        self.generate()
        self.generate({"hint": "Décrire l'origine des données"})
        self.assertIn("[Décrire l'origine des données]", self.texts())
        self.assertNotIn("[À compléter]", self.texts())

    def test_hint_ajoute_ne_deloge_pas_une_zone_redigee(self):
        self.generate()
        self.rewrite("[À compléter]", "Les données viennent du SI commercial.")
        self.generate({"hint": "Décrire l'origine des données"})
        self.assertIn("Les données viennent du SI commercial.", self.texts())
        self.assertNotIn("[Décrire l'origine des données]", self.texts())


if __name__ == "__main__":
    unittest.main()
