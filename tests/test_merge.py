"""Tests des marqueurs, du découpage en blocs et du bilan de mise à jour."""

import os
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from src import console
from src.merge import CHANGED, NEW, UNCHANGED, ChangeLog, markers, read_previous, salvage
from src.merge import blocks as block_parser
from src.merge.blocks import FREE, OWNED
from src.merge.previous import PreviousDocument


class MarkerFormatTest(unittest.TestCase):
    def test_element_aller_retour(self):
        marker = markers.parse(markers.element("measure:Marge", "abc123"))
        self.assertEqual(
            (marker.kind, marker.value, marker.fingerprint), ("elem", "measure:Marge", "abc123")
        )

    def test_identifiant_contenant_des_deux_points(self):
        self.assertEqual(markers.parse(markers.element("visual:p:v", "d1")).value, "visual:p:v")

    def test_contenu_genere(self):
        self.assertEqual(
            markers.parse(markers.opening(markers.GENERATED, "mesure_code")).value, "mesure_code"
        )
        self.assertEqual(markers.parse(markers.closing(markers.GENERATED)).kind, "endgen")

    def test_texte_ordinaire_non_reconnu(self):
        self.assertIsNone(markers.parse("Chiffre d'affaires"))
        self.assertIsNone(markers.parse(""))
        self.assertIsNone(markers.parse("pbi::inconnu|x"))

    def test_empreinte_ignore_la_mise_en_forme(self):
        self.assertEqual(
            markers.fingerprint("DIVIDE([A],\n   [B])"), markers.fingerprint("DIVIDE([A], [B])")
        )

    def test_empreinte_distingue_deux_expressions(self):
        self.assertNotEqual(markers.fingerprint("[A] + 1"), markers.fingerprint("[A] + 2"))

    def test_marqueur_ecrit_masque(self):
        doc = Document()
        markers.write(doc, markers.element("measure:CA", "d1"))
        self.assertIsNotNone(doc.paragraphs[-1].runs[0]._r.rPr.find(qn("w:vanish")))

    def test_marqueur_sans_place_dans_la_page(self):
        """La marque de paragraphe masquée, c'est la ligne du marqueur en moins."""
        doc = Document()
        markers.write(doc, markers.element("measure:CA", "d1"))

        properties = doc.paragraphs[-1]._p.pPr
        mark = properties.find(qn("w:rPr"))
        self.assertIsNotNone(mark.find(qn("w:vanish")))

        spacing = properties.find(qn("w:spacing"))
        self.assertEqual(spacing.get(qn("w:before")), "0")
        self.assertEqual(spacing.get(qn("w:after")), "0")
        self.assertEqual(spacing.get(qn("w:lineRule")), "exact")

    def test_marqueur_ecrit_en_petit(self):
        """Affiché, un marqueur reste lisible sans écarter ce qu'il encadre."""
        doc = Document()
        markers.write(doc, markers.closing(markers.GENERATED))

        run = doc.paragraphs[-1].runs[0]
        self.assertEqual(run.font.size, Pt(5))
        self.assertEqual(doc.paragraphs[-1]._p.pPr.find(qn("w:spacing")).get(qn("w:line")), "120")

    def test_marqueurs_d_une_version_anterieure_resserres(self):
        """Un marqueur recopié d'un ancien document est resserré à son tour."""
        doc = Document()
        paragraph = doc.add_paragraph()
        markers.hide(paragraph.add_run(markers.element("measure:CA", "d1")))
        libre = doc.add_paragraph("Texte de l'utilisateur")

        markers.collapse_all(doc)

        self.assertIsNotNone(paragraph._p.pPr.find(qn("w:rPr")).find(qn("w:vanish")))
        self.assertIsNone(libre._p.pPr)

    def test_marqueur_relu_depuis_le_xml(self):
        doc = Document()
        markers.write(doc, markers.opening(markers.GENERATED, "mesure_code"))
        self.assertEqual(markers.of(doc.paragraphs[-1]._p).value, "mesure_code")

    def test_paragraphe_ordinaire_sans_marqueur(self):
        doc = Document()
        doc.add_paragraph("Texte libre")
        self.assertIsNone(markers.of(doc.paragraphs[-1]._p))


class BlockParsingTest(unittest.TestCase):
    def _document(self):
        doc = Document()
        doc.add_paragraph("Avant toute ancre")
        markers.write(doc, markers.element("measure:Marge", "empreinte1"))
        doc.add_paragraph("Marge — titre reformulé")
        doc.add_paragraph("Note de l'utilisateur")
        markers.write(doc, markers.opening(markers.GENERATED, "mesure_code"))
        doc.add_paragraph("DIVIDE([A], [B])")
        markers.write(doc, markers.closing(markers.GENERATED))
        doc.add_paragraph("Explication libre")
        markers.write(doc, markers.element("measure:CA", "empreinte2"))
        doc.add_paragraph("CA")
        return doc

    def setUp(self):
        self.blocks = block_parser.parse(block_parser.body_nodes(self._document()))

    def test_un_bloc_par_ancre_plus_l_entete(self):
        self.assertEqual([b.element_id for b in self.blocks], ["", "measure:Marge", "measure:CA"])

    def test_empreinte_portee_par_le_bloc(self):
        self.assertEqual(self.blocks[1].fingerprint, "empreinte1")

    def test_contenu_du_script_isole(self):
        self.assertEqual(self.blocks[1].block_ids, ["mesure_code"])

    def test_tout_le_reste_appartient_a_l_utilisateur(self):
        free = [
            "".join(t.text or "" for t in node.iter(qn("w:t")))
            for node in self.blocks[1].free_nodes()
        ]
        self.assertEqual(
            free, ["Marge — titre reformulé", "Note de l'utilisateur", "Explication libre"]
        )

    def test_segment_du_script_porte_ses_delimiteurs(self):
        owned = next(s for s in self.blocks[1].segments if s.kind == OWNED)
        # marqueur d'ouverture + contenu + marqueur de fermeture
        self.assertEqual(len(owned.nodes), 3)

    def test_ordre_des_segments_conserve(self):
        self.assertEqual([s.kind for s in self.blocks[1].segments], [FREE, OWNED, FREE])

    def test_index_ecarte_le_contenu_hors_ancre(self):
        self.assertEqual(sorted(block_parser.index(self.blocks)), ["measure:CA", "measure:Marge"])


class SalvageTest(unittest.TestCase):
    """Ce que l'utilisateur a écrit à l'intérieur d'un contenu du script."""

    # Ce que le script avait écrit : un sous-titre, une valeur, une ligne vide.
    WRITTEN = ("Champs", "Ventes[Montant]", "")

    def _digests(self) -> list[str]:
        doc = Document()
        for text in self.WRITTEN:
            doc.add_paragraph(text)
        return [markers.digest(paragraph._p) for paragraph in doc.paragraphs]

    def _segment(self, *contents: str, forget: bool = False):
        """Segment relu dans le document, `contents` étant ce qu'on y trouve."""
        doc = Document()
        markers.write(doc, markers.opening(markers.GENERATED, "champs"))
        for text in contents:
            if text == "[tableau]":
                doc.add_table(rows=1, cols=1).cell(0, 0).text = "Ventes[Montant]"
            else:
                doc.add_paragraph(text)
        # Un document produit par une version antérieure : marqueur de fin nu.
        markers.write(
            doc, "pbi::endgen" if forget else markers.closing(markers.GENERATED, self._digests())
        )

        blocks = block_parser.parse(block_parser.body_nodes(doc))
        return blocks[0].segments[0]

    def _salvaged(self, segment) -> list[tuple[int, str]]:
        return [
            (rank, "".join(t.text or "" for t in node.iter(qn("w:t"))))
            for rank, node in salvage.of(segment)
        ]

    def test_rien_a_recuperer_quand_le_script_se_retrouve(self):
        self.assertEqual(self._salvaged(self._segment(*self.WRITTEN)), [])

    def test_texte_ecrit_dans_la_ligne_laissee_vide(self):
        segment = self._segment("Champs", "Ventes[Montant]", "Lecture du tableau.")
        self.assertEqual(self._salvaged(segment), [(2, "Lecture du tableau.")])

    def test_paragraphe_ajoute_au_milieu(self):
        segment = self._segment("Champs", "Une remarque.", "Ventes[Montant]", "")
        self.assertEqual(self._salvaged(segment), [(1, "Une remarque.")])

    def test_donnee_du_script_remaniee_a_la_main_abandonnee(self):
        """Le script reste propriétaire de ses données, même retouchées."""
        self.assertEqual(self._salvaged(self._segment("Champs", "Ventes[Autre]", "")), [])

    def test_donnee_remaniee_et_note_ajoutee(self):
        segment = self._segment("Champs", "Ventes[Autre]", "Ma note.", "")
        self.assertEqual(self._salvaged(segment), [(1, "Ma note.")])

    def test_valeur_disparue_du_script_abandonnee(self):
        self.assertEqual(self._salvaged(self._segment("Champs", "")), [])

    def test_ligne_vide_ajoutee_ignoree(self):
        segment = self._segment("Champs", "Ventes[Montant]", "", "")
        self.assertEqual(self._salvaged(segment), [])

    def test_version_anterieure_sans_tableau_ne_devine_rien(self):
        segment = self._segment("Champs", "Ventes[Montant]", "Ma note.", forget=True)
        self.assertEqual(self._salvaged(segment), [])

    def test_version_anterieure_rend_ce_qui_suit_le_tableau(self):
        segment = self._segment("Champs", "[tableau]", "Lecture du tableau.", forget=True)
        self.assertEqual(self._salvaged(segment), [(2, "Lecture du tableau.")])


class ReadPreviousTest(unittest.TestCase):
    def _read(self, doc):
        directory = tempfile.mkdtemp()
        path = os.path.join(directory, "doc.docx")
        doc.save(path)
        with console.silenced():
            return read_previous(path)

    def test_empreintes_relues(self):
        doc = Document()
        markers.write(doc, markers.element("measure:Marge", "d1"))
        doc.add_paragraph("Marge")
        previous = self._read(doc)
        self.assertTrue(previous.exists)
        self.assertEqual(previous.fingerprints, {"measure:Marge": "d1"})

    def test_document_absent(self):
        self.assertFalse(read_previous("/introuvable.docx").exists)

    def test_document_sans_marqueur_ignore(self):
        doc = Document()
        doc.add_paragraph("Documentation rédigée à la main")
        self.assertFalse(self._read(doc).exists)

    def test_document_reste_ouvert_pour_la_fusion(self):
        doc = Document()
        markers.write(doc, markers.element("measure:CA", "d1"))
        self.assertIsNotNone(self._read(doc).document)


class StatusTest(unittest.TestCase):
    def setUp(self):
        self.previous = PreviousDocument(
            path="doc.docx", fingerprints={"measure:Marge": "d1", "measure:CA": "d2"}
        )

    def test_element_inchange(self):
        self.assertEqual(self.previous.status("measure:Marge", "d1"), UNCHANGED)

    def test_element_modifie(self):
        self.assertEqual(self.previous.status("measure:Marge", "d9"), CHANGED)

    def test_element_nouveau(self):
        self.assertEqual(self.previous.status("measure:Panier", "d3"), NEW)

    def test_premiere_generation_tout_inchange(self):
        self.assertEqual(PreviousDocument().status("measure:Marge", "d1"), UNCHANGED)

    def test_elements_retires(self):
        self.assertEqual(self.previous.removed({"measure:Marge"}), ["measure:CA"])


class ChangeLogTest(unittest.TestCase):
    def _log(self, **kwargs) -> ChangeLog:
        log = ChangeLog(is_update=True)
        for status, names in kwargs.items():
            for name in names:
                log.record(name, {"new": NEW, "changed": CHANGED, "unchanged": UNCHANGED}[status])
        return log

    def test_premiere_generation(self):
        log = ChangeLog(is_update=False)
        log.record("measure:CA", UNCHANGED)
        self.assertIn("Première génération", log.summary())

    def test_aucun_changement(self):
        log = self._log(unchanged=["measure:CA"])
        self.assertFalse(log.has_changes)
        self.assertIn("Aucun changement", log.summary())

    def test_resume_des_changements(self):
        log = self._log(new=["measure:Panier"], changed=["measure:Marge"])
        log.removed = ["visual:p1:v2"]
        summary = log.summary()
        self.assertIn("1 élément(s) ajouté(s)", summary)
        self.assertIn("1 élément(s) modifié(s)", summary)
        self.assertIn("1 élément(s) retiré(s)", summary)

    def test_etat_par_element(self):
        log = self._log(new=["a"], changed=["b"], unchanged=["c"])
        self.assertEqual(log.status_of("b"), CHANGED)
        self.assertEqual(log.status_of("inconnu"), UNCHANGED)
        self.assertEqual(log.written_ids, {"a", "b", "c"})

    def test_details(self):
        log = self._log(new=["measure:Panier"])
        log.preserved = 3
        details = "\n".join(log.details())
        self.assertIn("measure:Panier", details)
        self.assertIn("3 contenu(s)", details)


if __name__ == "__main__":
    unittest.main()
