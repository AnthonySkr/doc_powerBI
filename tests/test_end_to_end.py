"""
Le parcours complet, sur le plan réellement livré.

Les autres tests vérifient le moteur avec des plans minimaux écrits pour
l'occasion. Celui-ci part d'un projet `.pbip` sur disque et de
`config_doc_pbi.yaml` tel qu'il est distribué : c'est le seul endroit où le
plan livré est lui-même mis à l'épreuve.

Ce qu'il verrouille, au-delà du fait que la génération aboutit : **rédiger une
zone laissée vide et régénérer conserve ce qui a été écrit**. Un bloc du plan
qui appartiendrait au script — un `property` et son repli — enverrait cette
rédaction en annexe à chaque génération, sans qu'aucun test unitaire ne puisse
le voir : le moteur, lui, ferait exactement son travail.
"""

import os
import shutil
import tempfile
import unittest

from docx import Document

from src import console
from src.cli.arguments import Options
from src.config import DEFAULT_CONFIG_PATH
from src.pipeline import run

FIXTURE = os.path.join(os.path.dirname(__file__), "fixtures", "rapport_test")

# Sous-titre de la zone visée, et style que le plan donne aux zones à rédiger.
_DESCRIPTION = "Description"
_TO_FILL_STYLE = "A completer"
_ORPHANS = "Contenu non replacé"


class EndToEndTest(unittest.TestCase):
    """Génération, rédaction dans le document, régénération."""

    @classmethod
    def setUpClass(cls):
        # Le projet est recopié : la génération écrit à côté du .pbip.
        cls._temp = tempfile.mkdtemp()
        cls.project = os.path.join(cls._temp, "rapport_test")
        shutil.copytree(FIXTURE, cls.project)
        cls.pbip = os.path.join(cls.project, "Rapport.pbip")

        with console.silenced():
            cls.output_dir = run(cls._options())
        cls.document = os.path.join(cls.output_dir, "documentation_Rapport.docx")

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls._temp, ignore_errors=True)

    @classmethod
    def _options(cls) -> Options:
        return Options(
            pbip_path=cls.pbip,
            config_path=DEFAULT_CONFIG_PATH,
            interactive=False,
            pause=False,
        )

    @classmethod
    def _regenerate(cls) -> None:
        with console.silenced():
            run(cls._options())

    def _paragraphs(self) -> list[str]:
        return [p.text.strip() for p in Document(self.document).paragraphs]

    # ── Génération ────────────────────────────────────────────────
    def test_le_document_est_produit(self):
        self.assertTrue(os.path.isfile(self.document))

    def test_le_contenu_du_rapport_est_documente(self):
        texts = self._paragraphs()
        for expected in ("Chiffre d'affaires", "CA N-1", "Synthèse commerciale", "Ventes"):
            self.assertIn(expected, texts, f"'{expected}' absent du document")

    def test_la_hierarchie_de_dates_tient_sur_une_reference(self):
        """L'axe du visuel porte une hiérarchie de dates : une ligne, pas trois."""
        cells = [
            cell.text.strip()
            for table in Document(self.document).tables
            for row in table.rows
            for cell in row.cells
        ]
        self.assertIn("Date (Année > Trimestre > Mois)", cells)
        for level in ("Date Année", "Date Trimestre", "Date Mois"):
            self.assertNotIn(level, cells, f"'{level}' occupe encore sa propre ligne")

    def test_le_code_dax_est_ecrit(self):
        self.assertTrue(
            any("SUM(Ventes[Montant])" in text for text in self._paragraphs()),
            "la formule DAX de la mesure n'apparaît pas",
        )

    def test_les_marqueurs_sont_poses(self):
        markers = [text for text in self._paragraphs() if text.startswith("pbi::")]
        self.assertTrue(markers, "aucun marqueur : la régénération repartirait de zéro")

    # ── Régénération ──────────────────────────────────────────────
    def test_la_description_dune_mesure_survit_a_la_regeneration(self):
        """
        Le contrat central, sur la zone que l'on remplit le plus.

        La cible est délibérément la description d'une mesure, et non la
        première zone venue : les parties `seed:` reviennent de toute façon
        intactes, et le test passerait sans rien prouver.
        """
        written = self._fill_description("Ce que cette mesure calcule, écrit à la main.")
        self._regenerate()

        texts = self._paragraphs()
        self.assertFalse(
            any(_ORPHANS in text for text in texts),
            "la rédaction est partie en annexe : le bloc du plan appartient au "
            "script alors qu'il invite à écrire",
        )
        self.assertIn(written, texts, "la rédaction n'a pas été conservée")

    def _fill_description(self, text: str) -> str:
        """
        Écrit la description d'une mesure, comme le ferait l'utilisateur.

        Repérée par son libellé plutôt que par son rang : c'est la zone qui
        suit le sous-titre « Description ».
        """
        document = Document(self.document)
        paragraphs = document.paragraphs
        for index, paragraph in enumerate(paragraphs[:-1]):
            if paragraph.text.strip() != _DESCRIPTION:
                continue
            zone = paragraphs[index + 1]
            if zone.style.name != _TO_FILL_STYLE or not zone.text.strip().startswith("["):
                continue
            zone.runs[0].text = text
            for extra in zone.runs[1:]:
                extra.text = ""
            document.save(self.document)
            return text
        self.fail("aucune description de mesure à rédiger dans le document généré")


if __name__ == "__main__":
    unittest.main()
