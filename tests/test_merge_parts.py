"""
Ce dont un paragraphe recopié dépend ailleurs dans le fichier.

Un paragraphe ne se suffit pas à lui-même : sa puce vit dans `numbering.xml`,
son style dans `styles.xml`, son commentaire dans `comments.xml`. Le document
neuf part du template et ne connaît rien de tout cela.
"""

import unittest

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from tests.test_merge_cycle import MergeHarness


def numbering(document):
    return document.part.part_related_by(RT.NUMBERING).element


def numbering_ids(document) -> set[str]:
    return {node.get(qn("w:numId")) for node in numbering(document).iter(qn("w:num"))}


def style_ids(document) -> set[str]:
    return {node.get(qn("w:styleId")) for node in document.styles.element.iter(qn("w:style"))}


class TransplantedPartsTest(MergeHarness):
    # ── Listes ────────────────────────────────────────────────────
    def test_liste_creee_dans_word_garde_sa_numerotation(self):
        self.generate(marge="SUM(T[a])")
        self._make_list("[À compléter]", "Premier point de ma liste", num_id="42")

        self.generate(marge="SUM(T[a])")

        document = Document(self.path)
        paragraph = next(p for p in document.paragraphs if "Premier point" in p.text)
        used = {node.get(qn("w:val")) for node in paragraph._p.iter(qn("w:numId"))}
        self.assertTrue(used <= numbering_ids(document), f"{used} absent de numbering.xml")

    def test_puce_definie_avec_sa_numerotation_abstraite(self):
        self.generate(marge="SUM(T[a])")
        self._make_list("[À compléter]", "Premier point", num_id="42")

        self.generate(marge="SUM(T[a])")

        document = Document(self.path)
        paragraph = next(p for p in document.paragraphs if "Premier point" in p.text)
        num_id = next(node.get(qn("w:val")) for node in paragraph._p.iter(qn("w:numId")))
        root = numbering(document)
        instance = next(n for n in root.iter(qn("w:num")) if n.get(qn("w:numId")) == num_id)
        abstract = instance.find(qn("w:abstractNumId")).get(qn("w:val"))
        self.assertIn(
            abstract,
            {n.get(qn("w:abstractNumId")) for n in root.iter(qn("w:abstractNum"))},
        )

    def test_numerotation_non_dupliquee_a_chaque_generation(self):
        self.generate(marge="SUM(T[a])")
        self._make_list("[À compléter]", "Premier point", num_id="42")
        self.generate(marge="SUM(T[a])")
        first = len(numbering_ids(Document(self.path)))

        self.generate(marge="SUM(T[a])")

        self.assertEqual(len(numbering_ids(Document(self.path))), first)

    # ── Styles ────────────────────────────────────────────────────
    def test_style_cree_dans_le_document_est_emmene(self):
        self.generate(marge="SUM(T[a])")
        document, paragraph = self.find("[À compléter]")
        style = document.styles.add_style("MonStylePerso", 1)
        style.font.name = "Consolas"
        self.clear_runs(paragraph)
        paragraph.add_run("Texte dans mon style")
        paragraph.style = style
        document.save(self.path)

        self.generate(marge="SUM(T[a])")

        document = Document(self.path)
        paragraph = next(p for p in document.paragraphs if "Texte dans mon style" in p.text)
        self.assertEqual(paragraph.style.name, "MonStylePerso")
        self.assertIn("MonStylePerso", style_ids(document))

    # ── Commentaires ──────────────────────────────────────────────
    def test_commentaire_word_conserve(self):
        """Un commentaire de révision suit le texte qu'il annote."""
        self.generate(marge="SUM(T[a])")
        document, paragraph = self.find("[À compléter]")
        self.clear_runs(paragraph)
        paragraph.add_run("Texte à revoir")
        document.add_comment(paragraph.runs, "Vérifier avec le métier", author="Anthony")
        document.save(self.path)

        self.generate(marge="SUM(T[a])")

        document = Document(self.path)
        self.assertIn("Texte à revoir", self.texts())
        comments = [comment.text for comment in document.comments]
        self.assertIn("Vérifier avec le métier", comments)

    def test_reference_de_commentaire_jamais_orpheline(self):
        """Une référence sans `comments.xml` rend le fichier illisible pour Word."""
        self.generate(marge="SUM(T[a])")
        document, paragraph = self.find("[À compléter]")
        self.clear_runs(paragraph)
        paragraph.add_run("Texte annoté")
        run = OxmlElement("w:r")
        reference = OxmlElement("w:commentReference")
        reference.set(qn("w:id"), "1")
        run.append(reference)
        paragraph._p.append(run)
        document.save(self.path)

        self.generate(marge="SUM(T[a])")

        document = Document(self.path)
        self.assertIn("Texte annoté", self.texts())
        references = list(document.element.body.iter(qn("w:commentReference")))
        if references:
            self.assertIsNotNone(
                document.part.part_related_by(
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
                )
            )

    # ── Utilitaire ────────────────────────────────────────────────
    def _make_list(self, needle: str, text: str, num_id: str) -> None:
        """Transforme un paragraphe en liste, comme le bouton « puces » de Word."""
        document, paragraph = self.find(needle)
        self.clear_runs(paragraph)
        paragraph.add_run(text)

        root = document.part.numbering_part.element
        instance = OxmlElement("w:num")
        instance.set(qn("w:numId"), num_id)
        abstract = OxmlElement("w:abstractNumId")
        abstract.set(qn("w:val"), "0")
        instance.append(abstract)
        root.append(instance)

        properties = paragraph._p.get_or_add_pPr()
        number = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        reference = OxmlElement("w:numId")
        reference.set(qn("w:val"), num_id)
        number.append(level)
        number.append(reference)
        properties.append(number)
        document.save(self.path)


if __name__ == "__main__":
    unittest.main()
