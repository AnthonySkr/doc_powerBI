"""
Une partie déclarée `seed:` appartient entièrement à l'utilisateur.

Les parties rédigées de bout en bout — « Initialisation », « Acquisition des
données » — sont écrites une fois, sous-titres compris, puis laissées telles
quelles. L'utilisateur y ajoute, y retire et y déplace ce qu'il veut : rien
n'est repéré à l'intérieur, donc rien n'est réécrit.
"""

import unittest

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from src.merge import markers
from tests.test_merge_cycle import MergeHarness

PLAN = {
    "document": {"template": "", "cover": {}, "properties": {}},
    "rendering": {"page_break_before_heading_1": False, "links": {"enabled": False}},
    "sections": [
        {
            "id": "initialisation",
            "title": "Initialisation",
            "level": 1,
            "seed": True,
            "sections": [
                {
                    "title": "Contexte",
                    "level": 2,
                    "blocks": [
                        {"type": "user_fill", "id": "ctx", "show_placeholder": False},
                    ],
                },
                {"title": "Objet", "level": 2, "blocks": []},
            ],
        },
        {"id": "suite", "title": "Suite", "level": 1, "blocks": []},
    ],
}


class FreeformSectionTest(MergeHarness):
    def build(self, **expressions):
        return self.generate(PLAN, **expressions)

    def after(self, needle: str, text: str) -> None:
        """Écrit dans le paragraphe qui suit un titre, comme dans Word."""
        document, paragraph = self.find(needle)
        Paragraph(paragraph._p.getnext(), document).add_run(text)
        document.save(self.path)

    def headings(self) -> list[str]:
        document = Document(self.path)
        return [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]

    def remove(self, needle: str) -> None:
        document, paragraph = self.find(needle)
        paragraph._p.getparent().remove(paragraph._p)
        document.save(self.path)

    # ── Première génération ───────────────────────────────────────
    def test_titres_ecrits_a_la_premiere_generation(self):
        self.build()
        self.assertEqual(self.headings(), ["Initialisation", "Contexte", "Objet", "Suite"])

    def test_aucun_reperage_a_l_interieur(self):
        """Un seul repère : la partie elle-même. Tout le reste est libre."""
        self.build()
        document = Document(self.path)
        anchors = [
            marker.value
            for p in document.paragraphs
            if (marker := markers.parse(p.text)) is not None and marker.kind == markers.ELEMENT
        ]
        inner = [p.text for p in document.paragraphs if p.text.startswith("pbi::seed|")]
        self.assertEqual(anchors, ["section:initialisation", "section:suite"])
        self.assertEqual(inner, [])

    # ── Régénération ──────────────────────────────────────────────
    def test_texte_ecrit_sous_un_titre_conserve(self):
        self.build()
        self.after("Contexte", "Le rapport suit les ventes Europe.")

        self.build()

        self.assertIn("Le rapport suit les ventes Europe.", self.texts())

    def test_titre_reformule_conserve(self):
        self.build()
        self.rewrite("Objet", "Objet et périmètre fonctionnel")

        self.build()

        self.assertIn("Objet et périmètre fonctionnel", self.headings())
        self.assertNotIn("Objet", self.headings())

    def test_sous_titre_supprime_ne_revient_pas(self):
        self.build()
        self.remove("Contexte")

        self.build()

        self.assertNotIn("Contexte", self.headings())

    def test_sous_titre_ajoute_conserve(self):
        self.build()
        document, paragraph = self.find("Objet")
        added = document.add_paragraph("Contraintes", style="Heading 2")
        paragraph._p.addnext(added._p)
        document.save(self.path)

        self.build()

        self.assertIn("Contraintes", self.headings())

    def test_rien_ne_part_en_annexe(self):
        self.build()
        self.after("Contexte", "Ma rédaction")
        self.remove("Objet")

        self.build()

        self.assertEqual(self.annexe(), [])

    def test_document_stable(self):
        self.build()
        self.after("Contexte", "Ma rédaction")
        self.build()
        first = self.layout()

        self.build()

        self.assertEqual(self.layout(), first)

    def test_partie_suivante_intacte(self):
        """La liberté s'arrête à la partie : la suivante reste repérée."""
        self.build()
        document = Document(self.path)
        anchors = [
            markers.parse(p.text).value
            for p in document.paragraphs
            if p.text.startswith("pbi::elem|section:suite")
        ]
        self.assertEqual(anchors, ["section:suite"])


class FieldsTest(MergeHarness):
    """Un contenu que Word calcule n'est pas de la rédaction."""

    def test_table_des_matieres_non_recueillie(self):
        self.build_with_toc()

        self.generate(PLAN)

        self.assertEqual(self.annexe(), [])

    def build_with_toc(self) -> None:
        self.generate(PLAN)
        document = Document(self.path)
        # Word recalcule le sommaire à l'ouverture : son texte ne ressemble
        # jamais à celui que le script avait posé.
        sdt = document.element.body.makeelement(qn("w:sdt"), {})
        run = sdt.makeelement(qn("w:t"), {})
        run.text = "Initialisation\t3"
        sdt.append(run)
        document.element.body.insert(0, sdt)
        document.save(self.path)


if __name__ == "__main__":
    unittest.main()
