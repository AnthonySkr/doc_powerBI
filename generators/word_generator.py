"""Génération du document Word de documentation."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models.data_models import PowerBIReport, DaxMeasure
import re

# ==============================================================================
#  MAPPING DES STYLES WORD
# ==============================================================================

# Style pour le titre principal sur la page de garde
STYLE_TITRE_PRINCIPAL = "Heading"

# Styles de titres pour la structure
STYLE_PAGE_INTERCALAIRE = "Heading 1"
STYLE_TITRE_PAGE_RAPPORT = "Heading 2"
STYLE_TITRE_VISUEL = "Heading 3"
STYLE_TITRE_MESURE = "Heading 4"

# Styles de texte
STYLE_NORMAL = "Normal"
STYLE_DESCRIPTION = "Intense Quote"
STYLE_CODE_DAX = "Code DAX"
STYLE_PROPRIETE = "Normal"
STYLE_PUCE = "List Bullet"

# ==============================================================================
#  HELPER POUR LIENS INTERNES (BOOKMARKS & HYPERLINKS)
# ==============================================================================


def safe_bookmark_name(name: str) -> str:
    """Crée un nom de signet valide pour Word (alphanumérique, sans espaces)."""
    # Ne garde que les lettres, chiffres et underscores. Word est restrictif.
    return re.sub(r"\W+", "", name)


def add_hyperlink(paragraph, text: str, bookmark_name: str):
    """Ajoute un hyperlien interne à un paragraphe."""
    part = paragraph.part
    r_id = part.relate_to(
        "#" + bookmark_name,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    # Applique le style "Hyperlink" qui est intégré à Word
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_props.append(style)
    run.append(run_props)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, bookmark_name: str):
    """Ajoute un signet (bookmark) à un paragraphe."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), "0")
    bm_start.set(qn("w:name"), bookmark_name)
    paragraph._p.insert(0, bm_start)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), "0")
    paragraph._p.append(bm_end)


# ==============================================================================


def generate_word_documentation(
    report: PowerBIReport,
    all_measures: dict[str, DaxMeasure],
    output_path: str = "documentation_powerbi.docx",
    template_path: str = "../template-doc-pbib.docx",
) -> str:
    """Génère la documentation Word en utilisant le template Eiffage."""

    try:
        doc = Document(template_path)
    except Exception as e:
        return f"ERREUR : Impossible de charger le template '{template_path}'. Détails : {e}"

    # Mise à jour du titre principal du document
    doc.core_properties.title = f"Documentation Rapport Power BI : {report.name}"

    # --- Page de garde ---
    # On va remplacer le placeholder "Titre du document" sur la première page.
    # C'est plus robuste que de supprimer et recréer.
    # (Cette partie est optionnelle et dépend de la structure exacte du template)
    # for p in doc.paragraphs:
    #     if "Titre du document" in p.text:
    #         p.text = "" # Efface le placeholder
    #         p.add_run(doc.core_properties.title).bold = True # Ajoute le vrai titre
    #         break

    doc.add_page_break()

    # --- Table des matières ---
    doc.add_heading("Table des matières", level=1)
    doc.add_paragraph(
        "Après avoir généré le document, veuillez faire un clic droit sur la table des matières ci-dessous et choisir 'Mettre à jour les champs' > 'Mettre à jour toute la table'."
    )

    doc.add_page_break()

    # --- Documentation des Pages et Visuels ---
    doc.add_paragraph("Documentation du Rapport", style=STYLE_PAGE_INTERCALAIRE)
    # PARTIE 1 : Visuels
    for page in report.pages:
        doc.add_paragraph(f"Page : {page.display_name}", style=STYLE_TITRE_PAGE_RAPPORT)
        doc.add_paragraph()

        for visual in sorted(page.visuals, key=lambda v: v.title):
            doc.add_paragraph(visual.title, style=STYLE_TITRE_VISUEL)

            p_type = doc.add_paragraph(style=STYLE_NORMAL)
            p_type.add_run("Type de visuel : ").bold = True
            p_type.add_run(visual.visual_type)

            if visual.elements:
                p_elements = doc.add_paragraph(style=STYLE_NORMAL)
                p_elements.add_run("Éléments de données :").bold = True

                for e in sorted(
                    visual.elements, key=lambda x: (x.role, x.display_name)
                ):
                    p_item = doc.add_paragraph(style=STYLE_PUCE)

                    if e.type_category == "Mesure":
                        # Si mesure -> hyperlien
                        measure_name = e.query_ref.split(".")[-1]
                        bookmark = safe_bookmark_name(measure_name)
                        add_hyperlink(p_item, e.display_name, bookmark)
                        p_item.add_run(f" (Rôle : {e.role}, Type : Mesure)")
                    else:
                        p_item.text = f"{e.display_name} (Rôle : {e.role}, Type : {e.type_category})"

            if visual.filters:
                p_filtres = doc.add_paragraph(style=STYLE_NORMAL)
                p_filtres.add_run("Filtres appliqués : ").bold = True
                p_filtres.add_run(
                    " ; ".join(sorted(set(f.to_string() for f in visual.filters)))
                )

            doc.add_paragraph()

    # PARTIE 2 : Mesures
    if all_measures:
        doc.add_page_break()
        doc.add_paragraph("Dictionnaire des Mesures", style=STYLE_PAGE_INTERCALAIRE)

        organized = _organize_measures(all_measures.keys(), all_measures)

        for table in sorted(organized.keys()):
            doc.add_paragraph(f"Table : {table}", style=STYLE_TITRE_PAGE_RAPPORT)
            for folder in sorted(organized[table].keys()):
                if folder != "Racine":
                    doc.add_paragraph(f"Dossier : {folder}", style=STYLE_TITRE_VISUEL)

                for measure_name in sorted(organized[table][folder]):
                    m = all_measures[measure_name]

                    # Titre mesure
                    p_mesure = doc.add_paragraph(m.name, style=STYLE_TITRE_MESURE)
                    # Ajout signet
                    add_bookmark(p_mesure, safe_bookmark_name(m.name))

                    if m.description:
                        doc.add_paragraph(m.description, style=STYLE_DESCRIPTION)

                    if m.expression:
                        doc.add_paragraph(m.expression, style=STYLE_CODE_DAX)

                    doc.add_paragraph()

    try:
        doc.save(output_path)
        return f"Documentation Word générée : '{output_path}'"
    except Exception as e:
        return f"Erreur lors de la sauvegarde du document : {e}"


def _organize_measures(
    measures_keys: set[str], all_measures: dict[str, DaxMeasure]
) -> dict[str, dict[str, list[str]]]:
    organized: dict[str, dict[str, list[str]]] = {}
    for name in sorted(list(measures_keys)):
        measure = all_measures[name]
        organized.setdefault(measure.table_name, {}).setdefault(
            measure.display_folder, []
        ).append(name)
    return organized
