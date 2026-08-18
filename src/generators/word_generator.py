"""
Génération du document Word à partir du plan décrit dans config_doc_pbi.yaml.

Le générateur ne connaît aucune structure de document : il parcourt les
`sections` / `blocks` de la configuration et écrit ce qu'elles décrivent.

Les mentions de mesures rencontrées dans les textes écrits (libellés des
visuels, code DAX, sources utilisées, paragraphes) sont automatiquement
transformées en liens internes vers la définition de la mesure.
"""

import hashlib
import os
import re
import unicodedata
from collections.abc import Callable
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from src.doc_config import DocConfig, evaluate, render, render_list, resolve
from src.generators.measure_links import MeasureLinker, collect_measures
from src.models.data_models import DocLink

TextProvider = Callable[[dict[str, Any], str], str]

# Mise en forme appliquée aux liens si le style de caractère est introuvable.
_LINK_FALLBACK_COLOR = "0563C1"

# Longueur maximale d'un nom de signet accepté par Word.
_BOOKMARK_MAX_LENGTH = 40

# Un champ de table des matières commence par le mot-clé TOC.
_TOC_FIELD = re.compile(r"^\s*TOC\b")

# Éléments de settings.xml devant suivre `w:updateFields` (ordre du schéma).
_SETTINGS_AFTER_UPDATE_FIELDS = (
    "w:hdrShapeDefaults",
    "w:footnotePr",
    "w:endnotePr",
    "w:compat",
    "w:docVars",
    "w:rsids",
    "w:mathPr",
    "w:attachedSchema",
    "w:themeFontLang",
    "w:clrSchemeMapping",
    "w:shapeDefaults",
    "w:decimalSymbol",
    "w:listSeparator",
)


def generate_word_documentation(
    config: DocConfig,
    context: dict[str, Any],
    output_path: str,
    text_provider: TextProvider | None = None,
) -> str:
    """
    Écrit le document Word.

    Args:
        config: configuration chargée depuis config_doc_pbi.yaml
        context: données exposées au plan (report, model, inputs, styles)
        output_path: chemin du .docx généré
        text_provider: callback optionnel permettant à l'utilisateur de
            modifier les textes des blocs `editable`
    """
    try:
        doc, _, _ = build_document(config, context, text_provider)
    except OSError as e:
        template_path = render(config.document.get("template"), context)
        return f"ERREUR : Impossible de charger le template '{template_path}'. Détails : {e}"

    try:
        doc.save(output_path)
    except Exception as e:  # noqa: BLE001
        return f"Erreur lors de la sauvegarde du document : {e}"

    if (config.rendering.get("table_of_contents") or {}).get("update_with_word"):
        print(f"  {refresh_fields_with_word(output_path)}")

    return f"Documentation Word générée : '{output_path}'"


def build_document(
    config: DocConfig,
    context: dict[str, Any],
    text_provider: TextProvider | None = None,
) -> tuple[Any, dict[str, str], dict[str, str]]:
    """
    Construit le document en mémoire, sans l'écrire.

    Le mode « mise à jour » compare ce document fraîchement généré au fichier
    déjà présent, pour n'y reporter que les nouveautés et les changements.

    Returns:
        (document, signets des blocs suivis -> id du bloc,
         signets d'items -> valeur brute du signet)
    """
    template_path = render(config.document.get("template"), context)
    doc = Document(template_path)
    builder = _DocumentBuilder(doc, config, context, text_provider)
    builder.build()
    return doc, builder.tracked_blocks, builder.item_bookmarks


def refresh_fields_with_word(path: str) -> str:
    """
    Recalcule les champs du document en pilotant Word (Windows uniquement).

    Sans cela, la table des matières est simplement marquée « à recalculer » et
    Word s'en charge à l'ouverture du fichier. Cette étape n'est utile que pour
    obtenir un document déjà à jour sans l'ouvrir — elle exige Word installé et
    le paquet `pywin32`.
    """
    try:
        import win32com.client  # type: ignore[import-not-found]
    except ImportError:
        return (
            "Word non piloté (pywin32 absent) : la table des matières sera "
            "recalculée à l'ouverture du document"
        )

    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        document = word.Documents.Open(os.path.abspath(path))
        document.Fields.Update()
        for toc in document.TablesOfContents:
            toc.Update()
        document.Save()
        return "Table des matières recalculée par Word"
    except Exception as e:  # noqa: BLE001
        return f"Word n'a pas pu recalculer les champs ({e}) — mise à jour à l'ouverture"
    finally:
        try:
            if document is not None:
                document.Close(False)
            if word is not None:
                word.Quit()
        except Exception:  # noqa: BLE001, S110
            pass


# ==============================================================================
#  Construction du document
# ==============================================================================


class _DocumentBuilder:
    def __init__(
        self,
        doc: Document,  # type: ignore[valid-type]
        config: DocConfig,
        context: dict[str, Any],
        text_provider: TextProvider | None = None,
    ):
        self.doc = doc
        self.config = config
        self.context = context
        self.text_provider = text_provider

        self._available_styles = {s.name for s in doc.styles}
        self._style_ids = {s.name: s.style_id for s in doc.styles}
        self._character_styles = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.CHARACTER}
        self._missing_styles: set[str] = set()
        self._bookmark_id = 0
        self._figure_number = 0

        # Signets écrits et ancres visées, pour vérifier qu'aucun lien ne pointe
        # dans le vide une fois le document terminé.
        self._bookmarks: set[str] = set()
        self._anchors: dict[str, int] = {}

        # Signet de la section en cours : préfixe des signets de blocs suivis.
        self._current_item: str = ""
        # Signets posés autour des blocs suivis : nom de signet -> id du bloc.
        self.tracked_blocks: dict[str, str] = {}
        # Signets d'items : nom de signet -> valeur brute (« measure:CA »).
        self.item_bookmarks: dict[str, str] = {}

        self._update = self.config.rendering.get("update") or {}
        self._links = self.config.rendering["links"]
        self._auto = self._links.get("auto") or {}
        self.linker = self._build_linker()
        self._link_targets = set(self.linker.targets.values()) if self.linker else set()

        self._block_writers = {
            "paragraph": self._write_paragraph_block,
            "image": self._write_image_block,
            "user_fill": self._write_user_fill_block,
            "property": self._write_property_block,
            "table": self._write_table_block,
            "loop": self._write_loop_block,
        }

    # ── Point d'entrée ────────────────────────────────────────────
    def build(self) -> None:
        self._write_cover()
        self._write_properties()
        self._write_header_footer()
        for section in self.config.sections:
            self._write_section(section, self.context)
        self._update_table_of_contents()
        self._report_links()

    # ── Page de garde et propriétés ───────────────────────────────
    def _write_cover(self) -> None:
        cover = self.config.document.get("cover") or {}
        placeholder = render(cover.get("placeholder"), self.context)
        text = render(cover.get("text"), self.context)
        if not placeholder or not text:
            return

        for paragraph in self.doc.paragraphs:
            if placeholder in paragraph.text:
                for run in list(paragraph.runs):
                    run._element.getparent().remove(run._element)
                self._write_rich_text(
                    paragraph,
                    text,
                    self.context,
                    bold=bool(cover.get("bold", True)),
                    links=False,
                )
                return

    def _write_properties(self) -> None:
        properties = self.config.document.get("properties") or {}
        for key, value in properties.items():
            if hasattr(self.doc.core_properties, key):
                setattr(self.doc.core_properties, key, render(value, self.context))

    # ── En-têtes et pieds de page du template ─────────────────────
    def _write_header_footer(self) -> None:
        """
        Remplace dans les en-têtes / pieds de page du template les textes
        déclarés dans `document.header_footer.replacements`.
        """
        rules = (self.config.document.get("header_footer") or {}).get("replacements") or []
        if not rules:
            return

        replaced = 0
        for section in self.doc.sections:
            containers = {
                "header": (
                    section.header,
                    section.first_page_header,
                    section.even_page_header,
                ),
                "footer": (
                    section.footer,
                    section.first_page_footer,
                    section.even_page_footer,
                ),
            }
            for scope, parts in containers.items():
                for part in parts:
                    # Un en-tête « lié au précédent » n'a pas de contenu propre :
                    # y toucher créerait une partie vide dans le document.
                    if part is None or part.is_linked_to_previous:
                        continue
                    for rule in rules:
                        if rule.get("scope", "all") not in ("all", scope):
                            continue
                        placeholder = render(rule.get("placeholder"), self.context)
                        value = render(rule.get("text"), self.context)
                        if not placeholder or not value:
                            continue
                        replaced += _replace_in_part(part, placeholder, value)

        if replaced:
            print(f"  En-tête / pied de page : {replaced} texte(s) remplacé(s)")

    # ── Table des matières ────────────────────────────────────────
    def _update_table_of_contents(self) -> None:
        """
        Demande à Word de recalculer la table des matières du template.

        Les numéros de page dépendent de la mise en page : seul Word peut les
        calculer. Le champ TOC est donc marqué « à recalculer » (`w:dirty`),
        ce que Word applique à l'ouverture du document.
        """
        options = self.config.rendering.get("table_of_contents") or {}
        if not options.get("update", True):
            return

        levels = render(options.get("levels"), self.context)
        fields = _mark_toc_fields_dirty(self.doc.element.body, levels)

        if not fields:
            print("  Aucun champ de table des matières trouvé dans le template")
            return

        if options.get("update_all_fields"):
            _set_update_fields(self.doc.settings.element)

        detail = f" (niveaux {levels})" if levels else ""
        print(f"  Table des matières{detail} : recalculée à l'ouverture dans Word")

    # ── Sections ──────────────────────────────────────────────────
    def _write_section(self, section: dict[str, Any], context: dict[str, Any]) -> None:
        if section.get("generate") is False or section.get("source") == "template":
            return
        if not evaluate(section.get("when"), context):
            return

        if self._needs_page_break(section):
            self.doc.add_page_break()

        previous_item = self._current_item

        title = render(section.get("title"), context)
        if title:
            level = int(section.get("level", 1))
            paragraph = self.doc.add_paragraph(title, style=self._style(f"heading_{level}"))
            self._add_title_suffix(paragraph, section, context)
            bookmark = section.get("bookmark")
            if bookmark:
                self._current_item = render(bookmark, context)
                name = self._bookmark_for(self._current_item)
                if name:
                    self.item_bookmarks[name] = self._current_item
                self._add_bookmark(paragraph, self._current_item)

        for block in section.get("blocks") or []:
            self._write_block(block, context)

        for child in section.get("sections") or []:
            self._write_section(child, context)

        self._current_item = previous_item

    def _add_title_suffix(
        self, paragraph, section: dict[str, Any], context: dict[str, Any]
    ) -> None:
        """
        Complète un titre par une information technique discrète (type du
        visuel, identifiant...), dans un style de caractère dédié.
        """
        suffix = render(section.get("title_suffix"), context)
        if not suffix:
            return

        separator = section.get("title_suffix_separator", "   ")
        run = paragraph.add_run(f"{separator}{suffix}")
        style = self._character_style(section.get("title_suffix_style") or "technical_id")
        if style:
            run.style = style

    def _needs_page_break(self, section: dict[str, Any]) -> bool:
        if "page_break_before" in section:
            return bool(section["page_break_before"])
        return bool(self.config.rendering.get("page_break_before_heading_1")) and (
            int(section.get("level", 1)) == 1
        )

    # ── Blocs ─────────────────────────────────────────────────────
    def _write_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        if not evaluate(block.get("when"), context):
            return

        writer = self._block_writers.get(block.get("type", ""))
        if writer is None:
            print(f"  Type de bloc inconnu ignoré : '{block.get('type')}' ({block.get('id')})")
            return

        if not self._tracks(block):
            writer(block, context)
            return

        # Le bloc est repérable dans le document : la mise à jour saura le
        # retrouver pour le comparer, le remplacer ou le signaler.
        # `before` garde une référence sur chaque élément : sans elle, lxml
        # recycle ses objets Python et la comparaison devient fausse.
        body = self.doc.element.body
        before = list(body)
        writer(block, context)
        written = [child for child in body if not any(child is known for known in before)]
        name = self._wrap_bookmark(written, block_bookmark(self._current_item, block["id"]))
        if name:
            self.tracked_blocks[name] = block["id"]

    def _write_paragraph_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        text = render(block.get("text"), context)
        if block.get("editable") and self.text_provider is not None:
            text = self.text_provider(block, text)
        if not text:
            return
        style = self._style(block.get("style") or "normal")
        paragraph = self.doc.add_paragraph(style=style)
        self._write_rich_text(paragraph, text, context, links=self._links_allowed(block, style))

    def _write_image_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        options = self.config.rendering["image_placeholder"]
        description = render(block.get("description"), context)

        if options.get("numbering"):
            self._figure_number += 1

        text = str(options.get("text_format", "[IMAGE] {description}")).format(
            description=description, n=self._figure_number
        )
        self.doc.add_paragraph(text, style=self._style(block.get("style") or "image"))

        if options.get("show_caption"):
            caption = str(options.get("caption_format", "{description}")).format(
                description=description, n=self._figure_number
            )
            self.doc.add_paragraph(caption, style=self._style("caption"))

        if options.get("empty_paragraph_after"):
            self.doc.add_paragraph()

    def _write_user_fill_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        options = self.config.rendering["user_fill"]
        text = ""
        if options.get("show_placeholder"):
            text = render(
                block.get("placeholder_text") or options.get("placeholder_text"),
                context,
            )
        style = block.get("style") or options.get("style") or "todo"
        self.doc.add_paragraph(text, style=self._style(style))

    def _write_property_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        options = self.config.rendering["property"]

        label = render(block.get("label"), context)
        if label:
            self.doc.add_paragraph(
                label,
                style=self._style(block.get("label_style") or options.get("label_style")),
            )

        value_style = self._style(block.get("value_style") or options.get("value_style"))
        links = self._links_allowed(block, value_style)
        fallback = render(block.get("fallback"), context)
        fallback_style = self._style(
            block.get("fallback_style") or options.get("fallback_style") or "todo"
        )

        if "value_list" in block:
            values = self._resolve_value_list(block.get("value_list"), context)
            if values:
                for value in values:
                    paragraph = self.doc.add_paragraph(style=value_style)
                    self._write_value(paragraph, value, context, links=links)
            elif fallback:
                self.doc.add_paragraph(fallback, style=fallback_style)
        else:
            value = render(block.get("value"), context)
            if value:
                paragraph = self.doc.add_paragraph(style=value_style)
                self._write_rich_text(paragraph, value, context, links=links)
            elif fallback:
                self.doc.add_paragraph(fallback, style=fallback_style)

        if options.get("empty_paragraph_after"):
            self.doc.add_paragraph()

    def _resolve_value_list(self, expression: Any, context: dict[str, Any]) -> list[Any]:
        """
        Résout `value_list` en conservant les objets `DocLink`, qui portent
        leur propre cible de lien (« Utilisée dans » d'une mesure).
        """
        if isinstance(expression, str) and "{{" in expression:
            items = render_list_of_items(expression, context)
        elif isinstance(expression, (list, tuple, set)):
            items = list(expression)
        else:
            items = render_list(expression, context)

        return [item for item in items if isinstance(item, DocLink) or render(item, context)]

    def _write_value(
        self, paragraph, value: Any, context: dict[str, Any], links: bool = True
    ) -> None:
        """Écrit une valeur de liste : lien vers un signet, ou texte enrichi."""
        if isinstance(value, DocLink):
            target = self._bookmark_for(value.target)
            if links and self._links.get("enabled", True) and self._is_reachable(target):
                self._add_hyperlink(paragraph, value.text, target)
                return
            self._write_rich_text(paragraph, value.text, context, links=links)
            return

        self._write_rich_text(paragraph, render(value, context), context, links=links)

    def _write_table_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        columns = block.get("columns") or []
        rows = render_list_of_items(block.get("over"), context)
        if not columns or not rows:
            return

        table = self.doc.add_table(rows=0, cols=len(columns))
        self._apply_table_style(table, render(block.get("style"), context))
        table.autofit = bool(block.get("autofit", False))

        # Largeurs fixes : Word suit la grille du tableau, pas les cellules.
        if block.get("layout", "fixed") == "fixed":
            _set_fixed_layout(table, [_column_width(c) for c in columns])
        _set_table_look(table, first_row=bool(block.get("header")))

        vertical_align = block.get("vertical_align", "center")

        if block.get("header"):
            labels = block.get("header_labels") or [c.get("id", "") for c in columns]
            header_row = table.add_row()
            if block.get("repeat_header", True):
                _repeat_header_row(header_row)
            for cell, label, column in zip(header_row.cells, labels, columns):
                _set_vertical_align(cell, vertical_align)
                style = column.get("header_style") or block.get("header_style")
                paragraph = cell.paragraphs[0]
                if style:
                    paragraph.style = self._style(style)
                paragraph.add_run(render(label, context))

        item_name = block.get("item") or "item"
        for row_item in rows:
            row_context = {**context, item_name: row_item}
            row = table.add_row()
            if block.get("cant_split", True):
                _keep_row_together(row)
            for cell, column in zip(row.cells, columns):
                _set_vertical_align(cell, vertical_align)
                self._fill_table_cell(cell, column, row_context)

        self._set_column_widths(table, columns)
        self.doc.add_paragraph()

    def _fill_table_cell(self, cell, column: dict[str, Any], context: dict[str, Any]) -> None:
        text = render(column.get("value"), context)
        paragraph = cell.paragraphs[0]

        if column.get("style"):
            paragraph.style = self._style(column.get("style"))

        # Lien explicite déclaré dans le plan : il prime sur la détection
        # automatique lorsque son texte est présent dans la cellule.
        hyperlink = column.get("hyperlink") or {}
        link_text = render(hyperlink.get("text"), context) if hyperlink else ""

        target = self._bookmark_for(render(hyperlink.get("target"), context)) if hyperlink else ""

        if (
            hyperlink
            and self._links.get("enabled", True)
            and link_text
            and link_text in text
            and self._is_reachable(target)
            and evaluate(hyperlink.get("when"), context)
        ):
            before, _, after = text.partition(link_text)
            if before:
                self._write_rich_text(paragraph, before, context)
            self._add_hyperlink(paragraph, link_text, target)
            if after:
                self._write_rich_text(paragraph, after, context)
            return

        self._write_rich_text(paragraph, text, context, links=self._links_allowed(column, ""))

    def _write_loop_block(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        items = render_list_of_items(block.get("over"), context)
        if not items:
            return

        item_name = block.get("item") or "item"
        section = block.get("section")
        blocks = block.get("blocks")

        for item in items:
            item_context = {**context, item_name: item}
            if section:
                self._write_section(section, item_context)
            for inner in blocks or []:
                self._write_block(inner, item_context)

    # ── Utilitaires de mise en forme ──────────────────────────────
    def _style(self, key: str | None) -> str:
        """Résout une clé de style (ou un `{{ styles.x }}`) en style Word existant."""
        if not key:
            key = "normal"

        name = self._style_name(key)
        if name in self._available_styles:
            return name

        fallback = self.config.styles.get("fallback", "Normal")
        if name and name != fallback and name not in self._missing_styles:
            self._missing_styles.add(name)
            print(f"  Style '{name}' absent du template — remplacé par '{fallback}'")
        return fallback if fallback in self._available_styles else "Normal"

    def _style_name(self, key: str | None) -> str:
        """Nom du style Word correspondant à une clé de configuration."""
        if not key:
            return ""
        return render(key, self.context) if "{{" in str(key) else self.config.styles.get(key, key)

    def _character_style(self, key: str | None) -> str | None:
        """Style de caractère existant, ou None (un style de paragraphe ne convient pas)."""
        name = self._style_name(key)
        if name in self._character_styles:
            return name
        if name and name not in self._missing_styles:
            self._missing_styles.add(name)
            print(f"  Style de caractère '{name}' absent du template — ignoré")
        return None

    def _apply_table_style(self, table, style_name: str) -> None:
        if not style_name:
            return
        if style_name not in self._available_styles:
            if style_name not in self._missing_styles:
                self._missing_styles.add(style_name)
                print(f"  Style de tableau '{style_name}' absent du template — ignoré")
            return
        try:
            table.style = style_name
        except (KeyError, ValueError) as e:
            print(f"  Style de tableau '{style_name}' non applicable : {e}")

    def _set_column_widths(self, table, columns: list[dict[str, Any]]) -> None:
        for index, column in enumerate(columns):
            width = _column_width(column)
            if width is None:
                continue
            for row in table.rows:
                row.cells[index].width = Cm(width)

    def _write_rich_text(
        self,
        paragraph,
        text: str,
        context: dict[str, Any],
        bold: bool = False,
        links: bool = True,
    ) -> None:
        """Écrit un texte en gérant les retours à la ligne et les liens internes."""
        lines = str(text).split("\n")
        for index, line in enumerate(lines):
            for chunk, bookmark in self._segments(line, context, links):
                if bookmark:
                    self._add_hyperlink(paragraph, chunk, bookmark)
                elif chunk:
                    run = paragraph.add_run(chunk)
                    run.bold = bold
            if index < len(lines) - 1:
                paragraph.add_run().add_break()

    def _segments(
        self, line: str, context: dict[str, Any], links: bool
    ) -> list[tuple[str, str | None]]:
        """Découpe une ligne en segments liés / non liés."""
        if not line:
            return []
        if not links or self.linker is None or not self.linker.has_targets():
            return [(line, None)]
        return self.linker.split(line, skip_bookmark=self._self_bookmark(context))

    def _self_bookmark(self, context: dict[str, Any]) -> str | None:
        """Signet de la mesure en cours de documentation (pas d'auto-référence)."""
        if not self._auto.get("skip_self", True):
            return None
        measure = context.get("measure")
        name = getattr(measure, "name", None)
        if not name or self.linker is None:
            return None
        return self.linker.targets.get(name)

    def _is_reachable(self, bookmark: str) -> bool:
        """
        Un lien explicite n'est écrit que si sa cible est une mesure documentée.

        Sans ce contrôle, une mesure absente du modèle (visuel pointant vers un
        autre jeu de données, mesure supprimée...) produirait dans Word un lien
        « Le signet n'existe pas ». Quand la détection automatique est
        désactivée, le plan reste maître de ses cibles.
        """
        if not bookmark:
            return False
        if self.linker is None or not self.linker.has_targets():
            return True
        return bookmark in self._link_targets or bookmark in self._bookmarks

    def _links_allowed(self, node: dict[str, Any], style_name: str) -> bool:
        """Détermine si la détection automatique s'applique à ce contenu."""
        if not self._links.get("enabled", True) or not self._auto.get("enabled", True):
            return False
        if node.get("links") is False or node.get("auto_links") is False:
            return False
        if not self._auto.get(  # noqa: SIM103
            "in_code", True
        ) and style_name == self._style("code"):
            return False
        return True

    # ── Signets et liens internes ─────────────────────────────────
    def _build_linker(self) -> MeasureLinker | None:
        """
        Construit le répertoire « nom de mesure -> signet » à partir des mesures
        effectivement documentées dans le document.
        """
        if not self._links.get("enabled", True) or not self._auto.get("enabled", True):
            return None

        source = self._auto.get("source") or "model.tables_with_measures"
        target_template = self._auto.get("target") or "measure:{{ measure.name }}"

        targets: dict[str, str] = {}
        for measure in collect_measures(resolve(str(source).strip("{} "), self.context)):
            name = getattr(measure, "name", "")
            bookmark = self._bookmark_for(
                render(target_template, {**self.context, "measure": measure})
            )
            if name and bookmark:
                targets[name] = bookmark

        known = self.context.get("report")
        known_names = list(getattr(known, "all_measures", {}) or {})

        # Mesures que l'on ne veut jamais lier automatiquement (nom trop
        # courant, mesure technique...).
        excluded = {str(name) for name in self._auto.get("exclude") or []}
        for name in excluded:
            targets.pop(name, None)
        known_names = [name for name in known_names if name not in excluded]

        return MeasureLinker(
            targets=targets,
            known_names=known_names,
            case_sensitive=bool(self._auto.get("case_sensitive", False)),
            min_length=int(self._auto.get("min_length", 2)),
            first_occurrence_only=bool(self._auto.get("first_occurrence_only", False)),
        )

    def _bookmark_for(self, raw_name: str) -> str:
        """Nom de signet d'une cible (`measure:Chiffre d'affaires` → signet Word)."""
        prefix = self._links.get("bookmark_prefix", "") or ""
        return _bookmark_name(prefix + (raw_name or ""))

    def _tracks(self, block: dict[str, Any]) -> bool:
        """Un bloc suivi porte un signet, pour être comparé d'une exécution à l'autre."""
        if not self._update.get("enabled", True) or not block.get("id"):
            return False
        if not self._current_item:
            return False
        return bool(block.get("track") or block.get("review"))

    def _wrap_bookmark(self, elements: list[Any], raw_name: str) -> str:
        """Encadre une suite d'éléments par un signet (début avant, fin après)."""
        name = self._bookmark_for(raw_name)
        if not name or not elements or name in self._bookmarks:
            return ""

        self._bookmarks.add(name)
        self._bookmark_id += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), name)
        elements[0].addprevious(start)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        elements[-1].addnext(end)
        return name

    def _add_bookmark(self, paragraph, raw_name: str) -> None:
        name = self._bookmark_for(raw_name)
        if not name or name in self._bookmarks:
            return

        self._bookmarks.add(name)
        self._bookmark_id += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), name)

        # `w:pPr` doit rester le premier enfant du paragraphe.
        properties = paragraph._p.find(qn("w:pPr"))
        if properties is not None:
            properties.addnext(start)
        else:
            paragraph._p.insert(0, start)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        paragraph._p.append(end)

    def _add_hyperlink(self, paragraph, text: str, bookmark: str) -> None:
        if not bookmark:
            paragraph.add_run(text)
            return

        self._anchors[bookmark] = self._anchors.get(bookmark, 0) + 1

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark)

        run = OxmlElement("w:r")
        run.append(self._link_run_properties())

        text_element = OxmlElement("w:t")
        text_element.text = text
        text_element.set(qn("xml:space"), "preserve")
        run.append(text_element)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _link_run_properties(self):
        """
        Mise en forme du lien. `w:rStyle` attend l'identifiant du style, pas son
        nom : dans un template français « Hyperlink » a pour identifiant
        « Lienhypertexte ». Si le style est absent, la mise en forme est
        appliquée directement pour que le lien reste visible.
        """
        properties = OxmlElement("w:rPr")
        style_name = self._links.get("style", "Hyperlink")
        style_id = self._style_ids.get(style_name)

        if style_id:
            style = OxmlElement("w:rStyle")
            style.set(qn("w:val"), style_id)
            properties.append(style)
            return properties

        color = OxmlElement("w:color")
        color.set(qn("w:val"), _LINK_FALLBACK_COLOR)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(color)
        properties.append(underline)
        return properties

    def _report_links(self) -> None:
        """Bilan des liens internes, affiché en fin de génération."""
        if self.linker is None:
            print("  Liens internes désactivés")
            return

        total = sum(self._anchors.values())
        measures = len(self._bookmarks & self._link_targets)
        others = len(self._bookmarks) - measures
        print(
            f"  Liens internes : {total} — {measures} mesures et "
            f"{others} autres emplacements atteignables"
        )

        dangling = sorted(set(self._anchors) - self._bookmarks)
        if dangling:
            print(
                f"  {len(dangling)} lien(s) sans signet correspondant : {', '.join(dangling[:5])}"
            )

        if self.linker.unlinked:
            names = sorted(self.linker.unlinked)
            print(
                f"  {len(names)} mesure(s) mentionnée(s) mais non documentée(s) "
                f"(hors périmètre `data.measures`) : {', '.join(names[:5])}"
            )


# ==============================================================================
#  Helpers
# ==============================================================================


def render_list_of_items(expression: Any, context: dict[str, Any]) -> list[Any]:
    """Résout l'expression `over:` d'une boucle ou d'un tableau."""
    if not expression:
        return []
    if isinstance(expression, (list, tuple)):
        return list(expression)

    path = str(expression).strip()
    if path.startswith("{{"):
        path = path.strip("{} ")

    value = resolve(path, context)
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        return list(value)
    if isinstance(value, set):
        return sorted(value)
    if isinstance(value, dict):
        return list(value.values())
    return [value]


def _column_width(column: dict[str, Any]) -> float | None:
    """Largeur d'une colonne en centimètres (`width_cm`)."""
    width = column.get("width_cm")
    return float(width) if width else None


def _set_fixed_layout(table, widths: list[float | None]) -> None:
    """
    Fige la mise en page du tableau et reporte les largeurs sur la grille.

    Word se fie à `w:tblGrid` : sans lui, les largeurs de cellules sont
    ignorées et les colonnes se répartissent au prorata du contenu.
    """
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    for column, width in zip(grid.findall(qn("w:gridCol")), widths):
        if width is not None:
            column.set(qn("w:w"), str(int(round(Cm(width).twips))))  # noqa: RUF046


def _set_table_look(table, first_row: bool) -> None:
    """
    Active la mise en forme conditionnelle du style de tableau : ligne
    d'en-tête colorée et lignes alternées, sans bandes verticales.
    """
    properties = table._tbl.tblPr
    look = properties.find(qn("w:tblLook"))
    if look is None:
        look = OxmlElement("w:tblLook")
        properties.append(look)

    values = {
        "firstRow": "1" if first_row else "0",
        "lastRow": "0",
        "firstColumn": "0",
        "lastColumn": "0",
        "noHBand": "0",
        "noVBand": "1",
    }
    look.set(qn("w:val"), "0620" if first_row else "0420")
    for key, value in values.items():
        look.set(qn("w:" + key), value)


def _repeat_header_row(row) -> None:
    """Répète la ligne d'en-tête en haut de chaque page."""
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        properties.append(OxmlElement("w:tblHeader"))


def _keep_row_together(row) -> None:
    """Empêche une ligne d'être coupée par un saut de page."""
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _set_vertical_align(cell, alignment: str) -> None:
    """Alignement vertical du contenu d'une cellule."""
    if not alignment:
        return
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:vAlign"))
    if element is None:
        element = OxmlElement("w:vAlign")
        properties.append(element)
    element.set(qn("w:val"), alignment)


def _replace_in_part(part, placeholder: str, value: str) -> int:
    """Remplace un texte dans un en-tête ou un pied de page."""
    replaced = 0
    for paragraph in part.paragraphs:
        replaced += _replace_in_paragraph(paragraph, placeholder, value)
    for table in part.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced += _replace_in_paragraph(paragraph, placeholder, value)
    return replaced


def _replace_in_paragraph(paragraph, placeholder: str, value: str) -> int:
    """
    Remplace un texte dans un paragraphe sans toucher au reste de sa mise en
    forme : seuls les nœuds `w:t` sont réécrits, les tabulations et les images
    qui les entourent (logo, alignement à droite de l'en-tête) sont conservées.
    """
    nodes = paragraph._p.findall(".//" + qn("w:t"))
    if not nodes:
        return 0

    # Cas courant : le texte tient dans un seul nœud.
    replaced = 0
    for node in nodes:
        text = node.text or ""
        if placeholder in text:
            replaced += text.count(placeholder)
            _set_node_text(node, text.replace(placeholder, value))
    if replaced:
        return replaced

    # Le texte est réparti sur plusieurs nœuds : on le regroupe sur le premier.
    joined = "".join(node.text or "" for node in nodes)
    if placeholder not in joined:
        return 0

    _set_node_text(nodes[0], joined.replace(placeholder, value))
    for node in nodes[1:]:
        _set_node_text(node, "")
    return joined.count(placeholder)


def _set_node_text(node, text: str) -> None:
    node.text = text
    if text != text.strip():
        node.set(qn("xml:space"), "preserve")


def _mark_toc_fields_dirty(body, levels: str = "") -> int:
    r"""
    Marque les champs TOC « à recalculer » et, si `levels` est renseigné,
    ajuste les niveaux de titres repris (`\o "1-3"`).

    Retourne le nombre de champs traités.
    """
    fields = 0

    # Champ complet : w:fldChar(begin) ... w:instrText ... w:fldChar(end)
    last_begin = None
    for element in body.iter(qn("w:fldChar"), qn("w:instrText")):
        if element.tag == qn("w:fldChar"):
            if element.get(qn("w:fldCharType")) == "begin":
                last_begin = element
            continue

        if not _TOC_FIELD.match(element.text or ""):
            continue

        if levels:
            element.text = _set_toc_levels(element.text or "", levels)
        if last_begin is not None:
            last_begin.set(qn("w:dirty"), "true")
        fields += 1

    # Forme condensée : w:fldSimple w:instr="TOC ..."
    for element in body.iter(qn("w:fldSimple")):
        instruction = element.get(qn("w:instr")) or ""
        if not _TOC_FIELD.match(instruction):
            continue
        if levels:
            element.set(qn("w:instr"), _set_toc_levels(instruction, levels))
        element.set(qn("w:dirty"), "true")
        fields += 1

    return fields


def _set_toc_levels(instruction: str, levels: str) -> str:
    r"""
    Remplace les niveaux repris par le champ TOC (`\\o "1-2"` → `\\o "1-3"`).

    Le remplacement passe par une fonction : la chaîne contient un antislash,
    que `re.sub` interpréterait comme une séquence d'échappement.
    """
    return re.sub(r'\\o\s*"[^"]*"', lambda _: f'\\o "{levels}"', instruction)


def _set_update_fields(settings) -> None:
    """
    Ajoute `<w:updateFields w:val="true"/>` : Word recalcule alors tous les
    champs du document à son ouverture. L'élément doit respecter l'ordre du
    schéma, sans quoi Word considère le fichier comme corrompu.
    """
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return

    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")

    for tag in _SETTINGS_AFTER_UPDATE_FIELDS:
        successor = settings.find(qn(tag))
        if successor is not None:
            successor.addprevious(element)
            return

    settings.append(element)


def style_id(doc, config: DocConfig, key: str) -> str:
    """Identifiant du style Word correspondant à une clé de configuration."""
    name = config.styles.get(key, key)
    for style in doc.styles:
        if style.name == name:
            return style.style_id
    return ""


def heading_style_levels(doc, config: DocConfig) -> dict[str, int]:
    """
    Identifiant de style de titre -> niveau (« Titre2 » -> 2).

    Les titres du document portent l'identifiant du style, pas son nom : la
    relecture d'un document a besoin de la correspondance pour reconstituer
    la hiérarchie des sections.
    """
    ids = {style.name: style.style_id for style in doc.styles}
    levels: dict[str, int] = {}
    for key, name in config.styles.items():
        if not key.startswith("heading_"):
            continue
        style_id = ids.get(name)
        if style_id:
            levels[style_id] = int(key.split("_")[1])
    return levels


def block_bookmark(item_bookmark: str, block_id: str) -> str:
    """Nom brut du signet encadrant un bloc suivi d'un item."""
    return f"{item_bookmark}#{block_id}"


def _bookmark_name(name: str) -> str:
    """
    Nom de signet valide pour Word : lettres non accentuées, chiffres et
    underscores, 40 caractères maximum, ne commençant pas par un chiffre.

    Dès que le nettoyage perd de l'information (accents, espaces, ponctuation,
    longueur), une empreinte du nom d'origine est ajoutée : sans elle
    « Marge » et « Marge % » produiraient le même signet et les deux mesures
    partageraient la même cible.

    La fonction reste purement déterministe : la pose du signet et la
    résolution des liens qui le visent aboutissent au même nom.
    """
    raw = name or ""
    ascii_name = unicodedata.normalize("NFKD", raw).encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"\W+", "_", ascii_name).strip("_")
    if not cleaned:
        cleaned = "signet"
    if cleaned[0].isdigit():
        cleaned = f"_{cleaned}"

    if cleaned == raw and len(cleaned) <= _BOOKMARK_MAX_LENGTH:
        return cleaned

    digest = hashlib.md5(raw.encode("utf-8")).hexdigest()[:6]
    return f"{cleaned[: _BOOKMARK_MAX_LENGTH - len(digest) - 1]}_{digest}"
