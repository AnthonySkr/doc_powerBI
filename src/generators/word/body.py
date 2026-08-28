"""
Écriture des contenus à la fin du corps du document.

python-docx place chaque paragraphe et chaque tableau *devant* `w:sectPr`, les
propriétés de section qui ferment le corps. Pour trouver cette place, il
parcourt tous les enfants du corps — à chaque écriture. Le coût d'un contenu
croît donc avec ce qui a déjà été écrit, et celui d'un document avec le carré
de sa taille : 1,7 s pour 50 mesures, 37 s pour 200, et plus d'aboutissement du
tout au-delà.

`Body` retient `w:sectPr` une fois pour toutes et insère juste devant : la
place est la même, le parcours n'a plus lieu. Le document produit est
identique, à ceci près qu'il s'écrit en temps constant.

Les styles suivent le même principe. `paragraph.style = "Titre"` demande à
python-docx l'identifiant du style, qui compare le style au style par défaut du
template en reparcourant toute sa table — pour chaque paragraphe. Les
identifiants sont ici relevés une fois à l'ouverture.
"""

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.table import CT_Tbl
from docx.shared import Emu, Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

_SECTION = qn("w:sectPr")


class Body:
    """Ajoute paragraphes et tableaux à la fin du corps, en temps constant."""

    def __init__(self, doc):
        self.doc = doc
        self._body = doc.element.body
        self._section = self._body.find(_SECTION)

        # Nom de style -> identifiant. Le style par défaut du template ne
        # s'écrit pas : python-docx l'omet, et un `w:pStyle` en trop ferait
        # diverger le document de celui des versions précédentes.
        self._style_ids = {style.name: style.style_id for style in doc.styles}
        default = doc.styles.default(WD_STYLE_TYPE.PARAGRAPH)
        self._default_style_id = default.style_id if default is not None else None

    # ── Écriture ──────────────────────────────────────────────────
    def add_paragraph(self, text: str = "", style: str | None = None) -> Paragraph:
        """Ajoute un paragraphe à la fin du corps."""
        paragraph = Paragraph(self._place(OxmlElement("w:p")), self.doc)
        if text:
            paragraph.add_run(text)
        if style is not None:
            paragraph._p.style = self._style_id(style)
        return paragraph

    def add_table(self, rows: int, cols: int) -> Table:
        """Ajoute un tableau à la fin du corps, ses colonnes de largeur égale."""
        return Table(self._place(CT_Tbl.new_tbl(rows, cols, self._width())), self.doc)

    def add_page_break(self) -> Paragraph:
        """Ajoute un saut de page à la fin du corps."""
        paragraph = self.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return paragraph

    # ── Interne ───────────────────────────────────────────────────
    def _place(self, element):
        """Insère l'élément devant les propriétés de section, sans les chercher."""
        if self._section is not None:
            self._section.addprevious(element)
        else:
            self._body.append(element)
        return element

    def _style_id(self, name: str) -> str | None:
        """
        Identifiant du style, ou None quand il n'y a rien à écrire.

        None pour le style par défaut du template — python-docx ne l'écrit pas
        non plus — et pour un nom inconnu : mieux vaut un paragraphe au style
        par défaut qu'une génération interrompue.
        """
        style_id = self._style_ids.get(name)
        return None if style_id == self._default_style_id else style_id

    def _width(self):
        """Largeur utile de la page, entre les marges de la dernière section."""
        section = self.doc.sections[-1]
        page_width = section.page_width or Inches(8.5)
        left = section.left_margin or Inches(1)
        right = section.right_margin or Inches(1)
        return Emu(page_width - left - right)
