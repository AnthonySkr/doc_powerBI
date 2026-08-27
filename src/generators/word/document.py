"""
Écriture du document Word en parcourant le plan de `config_doc_pbi.yaml`.

Le builder ne connaît aucune structure de document : il déroule les `sections`
et les `blocks` de la configuration et écrit ce qu'elles décrivent, à la suite
du contenu déjà présent dans le template.
"""

import math
from collections.abc import Callable
from typing import Any

from docx.shared import Cm, Emu, Pt

from src import console
from src.config import DocConfig, evaluate, render, render_list, resolve_items
from src.generators.word import fields, shapes, tables
from src.generators.word.links import LinkIndex
from src.generators.word.merging import MergeWriter
from src.generators.word.styles import StyleResolver
from src.merge import PreviousDocument
from src.models.data_models import DocLink

# Callback proposant à l'utilisateur de réécrire le texte d'un bloc `editable`.
TextProvider = Callable[[dict[str, Any], str], str]


class DocumentBuilder:
    def __init__(
        self,
        doc,
        config: DocConfig,
        context: dict[str, Any],
        text_provider: TextProvider | None = None,
        previous: PreviousDocument | None = None,
    ):
        self.doc = doc
        self.config = config
        self.context = context
        self.text_provider = text_provider

        self.styles = StyleResolver(doc, config, context)
        self.links = LinkIndex(config, context, self.styles.ids)
        self.merge = MergeWriter(doc, config, previous)
        self._figure_number = 0
        # Word refuse deux formes de même identifiant : la numérotation des
        # repères reprend au-dessus de ce que le template contient déjà.
        self._shape_id = shapes.last_id(doc)

        self._block_writers = {
            "paragraph": self._write_paragraph,
            "image": self._write_image,
            "user_fill": self._write_user_fill,
            "property": self._write_property,
            "table": self._write_table,
        }

    # ── Point d'entrée ────────────────────────────────────────────
    def build(self) -> None:
        self._write_cover()
        self._write_properties()
        self._write_header_footer()
        for section in self.config.sections:
            self._write_section(section, self.context, "")
        self._update_table_of_contents()
        self.links.report()

    # ── Éléments du template ──────────────────────────────────────
    def _write_cover(self) -> None:
        """Remplace le texte repère de la page de garde par le titre du rapport."""
        cover = self.config.document.get("cover") or {}
        placeholder = render(cover.get("placeholder"), self.context)
        text = render(cover.get("text"), self.context)
        if not placeholder or not text:
            return

        for paragraph in self.doc.paragraphs:
            if placeholder not in paragraph.text:
                continue
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            self._write_rich_text(
                paragraph,
                text,
                self.context,
                bold=bool(cover.get("bold", True)),
                links=False,
            )
            return

    def _write_properties(self) -> None:
        """Renseigne les propriétés du fichier .docx (titre, sujet, catégorie...)."""
        for key, value in (self.config.document.get("properties") or {}).items():
            if hasattr(self.doc.core_properties, key):
                setattr(self.doc.core_properties, key, render(value, self.context))

    def _write_header_footer(self) -> None:
        """
        Remplace dans les en-têtes / pieds de page du template les textes
        déclarés dans `document.header_footer.replacements`.
        """
        rules = (self.config.document.get("header_footer") or {}).get("replacements") or []
        if not rules:
            return

        replaced = 0
        for section in self.doc.sections:
            for scope, parts in _header_footer_parts(section).items():
                for part in parts:
                    # Un en-tête « lié au précédent » n'a pas de contenu propre :
                    # y toucher créerait une partie vide dans le document.
                    if part is None or part.is_linked_to_previous:
                        continue
                    for rule in rules:
                        if rule.get("scope", "all") not in ("all", scope):
                            continue
                        placeholder = render(rule.get("placeholder"), self.context)
                        value = render(rule.get("text"), self.context)
                        if placeholder and value:
                            replaced += fields.replace_in_part(part, placeholder, value)

        if replaced:
            console.info(f"En-tête / pied de page : {replaced} texte(s) remplacé(s)")

    def _update_table_of_contents(self) -> None:
        """Marque la table des matières du template comme à recalculer."""
        options = self.config.rendering.get("table_of_contents") or {}
        if not options.get("update", True):
            return

        levels = render(options.get("levels"), self.context)
        if not fields.mark_toc_fields_dirty(self.doc.element.body, levels):
            console.warn("Aucun champ de table des matières trouvé dans le template")
            return

        if options.get("update_all_fields"):
            fields.set_update_fields(self.doc.settings.element)

        detail = f" (niveaux {levels})" if levels else ""
        console.info(f"Table des matières{detail} : recalculée à l'ouverture dans Word")

    # ── Sections ──────────────────────────────────────────────────
    def _write_section(
        self, section: dict[str, Any], context: dict[str, Any], parent: str = ""
    ) -> None:
        """
        Écrit une section du plan.

        `parent` est l'identifiant de la partie qui la contient : il sert à
        repérer une sous-partie que le plan n'identifie pas autrement (ni
        `bookmark:`, ni `id:`), et se transmet à ses propres filles.

        Une partie déclarée `seed:` fait exception : son contenu — sous-parties
        comprises — n'est pas repéré du tout, et lui revient en bloc.
        """
        if section.get("generate") is False or section.get("source") == "template":
            return
        if not evaluate(section.get("when"), context):
            return

        if self._needs_page_break(section):
            self.doc.add_page_break()

        # L'ancre précède le titre : à la relecture, tout ce qui suit lui
        # appartient jusqu'à l'ancre suivante.
        element_id = self.merge.anchor(section, context, parent) or parent
        self._write_title(section, context)

        # `seed:` — la partie est écrite une fois, puis appartient entièrement
        # à l'utilisateur : ses sous-titres et ses textes ne sont plus repérés,
        # et lui reviennent tels qu'il les a laissés.
        with self.merge.freeform(bool(section.get("seed"))):
            for block in section.get("blocks") or []:
                self._write_block(block, context, element_id)

            for child in section.get("sections") or []:
                self._write_section(child, context, element_id)

    def _write_title(self, section: dict[str, Any], context: dict[str, Any]) -> None:
        title = render(section.get("title"), context)
        if not title:
            return

        level = int(section.get("level", 1))
        paragraph = self.doc.add_paragraph(title, style=self.styles.paragraph(f"heading_{level}"))
        self._add_title_suffix(paragraph, section, context)
        if section.get("bookmark"):
            self.links.add_bookmark(paragraph, render(section["bookmark"], context))

    def _add_title_suffix(self, paragraph, section: dict[str, Any], ctx: dict[str, Any]) -> None:
        """Complète un titre par une mention technique discrète (type du visuel...)."""
        suffix = render(section.get("title_suffix"), ctx)
        if not suffix:
            return

        run = paragraph.add_run(f"{section.get('title_suffix_separator', '   ')}{suffix}")
        style = self.styles.character(section.get("title_suffix_style") or "technical_id")
        if style:
            run.style = style

    def _needs_page_break(self, section: dict[str, Any]) -> bool:
        if "page_break_before" in section:
            return bool(section["page_break_before"])
        return bool(self.config.rendering.get("page_break_before_heading_1")) and (
            int(section.get("level", 1)) == 1
        )

    # ── Blocs ─────────────────────────────────────────────────────
    def _write_block(
        self, block: dict[str, Any], context: dict[str, Any], parent: str = ""
    ) -> None:
        if not evaluate(block.get("when"), context):
            return

        if block.get("type") == "loop":
            # Une boucle n'écrit rien elle-même : elle déroule un sous-plan,
            # dont chaque section porte sa propre identité. Elle n'est donc
            # jamais encadrée.
            self._write_loop(block, context, parent)
            return

        writer = self._block_writers.get(block.get("type", ""))
        if writer is None:
            console.warn(f"Type de bloc inconnu ignoré : '{block.get('type')}' ({block.get('id')})")
            return

        with self.merge.delimit(block):
            writer(block, context)

    def _write_paragraph(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        text = render(block.get("text"), context)
        if block.get("editable") and self.text_provider is not None:
            text = self.text_provider(block, text)
        if not text:
            return

        style = self.styles.paragraph(block.get("style") or "normal")
        paragraph = self.doc.add_paragraph(style=style)
        self._write_rich_text(paragraph, text, context, links=self._links_allowed(block, style))

    def _write_image(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        """Réserve l'emplacement d'une capture, avec sa description."""
        options = self.config.rendering["image_placeholder"]
        description = render(block.get("description"), context)
        numbering = _numbering_mode(options.get("numbering"))

        if numbering != "none":
            self._figure_number += 1

        number = str(self._figure_number) if numbering != "none" else ""
        text = str(options.get("text_format", "[IMAGE] {description}")).format(
            description=description, n=number
        )
        self.doc.add_paragraph(text, style=self.styles.paragraph(block.get("style") or "image"))

        if options.get("show_caption"):
            self._write_caption(options, description, number, numbering)

        self._write_markers(block, context, options)

        if options.get("empty_paragraph_after"):
            self.doc.add_paragraph()

    def _write_caption(
        self, options: dict[str, Any], description: str, number: str, numbering: str
    ) -> None:
        """
        Légende numérotée de la capture.

        Le numéro est un champ Word (`SEQ`), pas un texte : supprimer une
        capture renumérote les suivantes à l'ouverture du document, sans
        reprise à la main. `numbering: fixed` le fige dans le texte, pour un
        document destiné à un lecteur qui ne recalcule pas les champs.
        """
        template = str(options.get("caption_format", "{description}"))
        values = {"description": description, "n": number}
        paragraph = self.doc.add_paragraph(style=self.styles.paragraph("caption"))

        head, field, tail = template.partition("{n}")
        if numbering != "auto" or not field:
            paragraph.add_run(template.format(**values))
            return

        paragraph.add_run(head.format(**values))
        fields.write_sequence_field(paragraph, str(options.get("sequence") or "Figure"), number)
        paragraph.add_run(tail.format(**values))

    def _write_markers(
        self, block: dict[str, Any], context: dict[str, Any], options: dict[str, Any]
    ) -> None:
        """
        Repères numérotés à faire glisser sur la capture.

        Les numéros sont ceux du tableau qui suit la capture — le plan désigne
        la même liste. Ils sont posés en rangée sous l'emplacement, et n'ont
        plus qu'à être déplacés un à un sur l'image : ce sont des formes
        flottantes, elles ne bousculent rien en route.
        """
        plan = block.get("markers")
        if not plan:
            return

        item = plan.get("item") or "item"
        labels = [
            render(
                plan.get("label") or f"{{{{ {item}.number }}}}",
                {**context, item: value},
            )
            for value in resolve_items(plan.get("over"), context)
        ]
        labels = [label for label in labels if label]
        if not labels:
            return

        look = options.get("markers") or {}
        size = Cm(float(look.get("size_cm", 0.62)))
        spacing = Cm(float(look.get("spacing_cm", 0.9)))
        line = Cm(float(look.get("line_cm", 0.9)))
        per_row = max(int(look.get("per_row", 12) or 1), 1)

        paragraph = self.doc.add_paragraph(style=self.styles.paragraph(look.get("style")))
        # Les repères flottent : sans hauteur réservée, ils déborderaient sur
        # le tableau qui suit. Le paragraphe porte donc celle de leurs rangées.
        paragraph.paragraph_format.line_spacing = Emu(int(line) * math.ceil(len(labels) / per_row))
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        positions = shapes.row_positions(len(labels), spacing, per_row, line)
        for label, (left, top) in zip(labels, positions):
            self._shape_id += 1
            paragraph._p.append(
                shapes.marker(
                    label,
                    self._shape_id,
                    left,
                    top,
                    size,  # type: ignore
                    shape=str(look.get("shape") or "ellipse"),
                    fill=str(look.get("fill") or "0070C0"),
                    text_color=str(look.get("text_color") or "FFFFFF"),
                    font_size=Pt(float(look.get("font_size_pt", 9))),
                )
            )

    def _write_user_fill(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        """
        Zone à rédiger après génération.

        Elle n'est écrite qu'à la première génération : ensuite, ce que
        l'utilisateur y a mis est repris tel quel par la fusion.
        """
        options = self.config.rendering["user_fill"]

        label = render(block.get("label"), context)
        if label:
            self.doc.add_paragraph(
                label,
                style=self.styles.paragraph(
                    block.get("label_style") or self.config.rendering["property"].get("label_style")
                ),
            )

        shown = block.get("show_placeholder", options.get("show_placeholder"))
        text = self._user_fill_text(block, context, options) if shown else ""
        style = block.get("style") or options.get("style") or "todo"
        self.doc.add_paragraph(text, style=self.styles.paragraph(style))

    def _user_fill_text(
        self, block: dict[str, Any], context: dict[str, Any], options: dict[str, Any]
    ) -> str:
        """
        Amorce d'une zone à rédiger.

        Le bloc du plan dit ce qu'on attend à cet endroit (`hint:`) : autant
        l'écrire, plutôt qu'un « [À compléter] » qui ne guide personne. Même
        mise en forme dans les deux cas — c'est la même invitation à écrire.
        """
        hint = render(block.get("hint"), context)
        if hint:
            return str(options.get("hint_format", "[{hint}]")).format(hint=hint)
        return render(block.get("placeholder_text") or options.get("placeholder_text"), context)

    def _write_property(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        """Sous-titre + valeur, ou sous-titre + liste de valeurs."""
        options = self.config.rendering["property"]

        label = render(block.get("label"), context)
        if label:
            self.doc.add_paragraph(
                label,
                style=self.styles.paragraph(block.get("label_style") or options.get("label_style")),
            )

        value_style = self.styles.paragraph(block.get("value_style") or options.get("value_style"))
        links = self._links_allowed(block, value_style)

        values = (
            self._value_list(block["value_list"], context)
            if "value_list" in block
            else [render(block.get("value"), context)]
        )
        values = [value for value in values if value]

        for value in values:
            paragraph = self.doc.add_paragraph(style=value_style)
            self._write_value(paragraph, value, context, links=links)

        fallback = render(block.get("fallback"), context)
        if not values and fallback:
            self.doc.add_paragraph(
                fallback,
                style=self.styles.paragraph(
                    block.get("fallback_style") or options.get("fallback_style") or "todo"
                ),
            )

        if options.get("empty_paragraph_after"):
            self.doc.add_paragraph()

    def _value_list(self, expression: Any, context: dict[str, Any]) -> list[Any]:
        """
        Résout `value_list` en conservant les objets `DocLink`, qui portent
        leur propre cible de lien (« Utilisée dans » d'une mesure).
        """
        if isinstance(expression, str) and "{{" in expression:
            items = resolve_items(expression, context)
        elif isinstance(expression, (list, tuple, set)):
            items = list(expression)
        else:
            items = render_list(expression, context)

        return [item for item in items if isinstance(item, DocLink) or render(item, context)]

    def _write_table(self, block: dict[str, Any], context: dict[str, Any]) -> None:
        columns = block.get("columns") or []
        rows = resolve_items(block.get("over"), context)
        if not columns or not rows:
            return

        # Sous-titre facultatif, comme pour un bloc `property` : il n'est écrit
        # que si le tableau l'est, et disparaît donc avec lui.
        label = render(block.get("label"), context)
        if label:
            self.doc.add_paragraph(
                label,
                style=self.styles.paragraph(
                    block.get("label_style") or self.config.rendering["property"].get("label_style")
                ),
            )

        table = self.doc.add_table(rows=0, cols=len(columns))
        style = self.styles.table(render(block.get("style"), context))
        if style:
            table.style = style
        table.autofit = bool(block.get("autofit", False))

        # Largeurs fixes : Word suit la grille du tableau, pas les cellules.
        if block.get("layout", "fixed") == "fixed":
            tables.set_fixed_layout(table, [_column_width(c) for c in columns])
        tables.set_table_look(table, first_row=bool(block.get("header")))

        vertical_align = block.get("vertical_align", "center")

        if block.get("header"):
            self._write_table_header(table, block, columns, context, vertical_align)

        item_name = block.get("item") or "item"
        for row_item in rows:
            row = table.add_row()
            if block.get("cant_split", True):
                tables.keep_row_together(row)
            for cell, column in zip(row.cells, columns):
                tables.set_vertical_align(cell, vertical_align)
                self._fill_cell(cell, column, {**context, item_name: row_item})

        _apply_column_widths(table, columns)
        self.doc.add_paragraph()

    def _write_table_header(
        self,
        table,
        block: dict[str, Any],
        columns: list[dict[str, Any]],
        context: dict[str, Any],
        vertical_align: str,
    ) -> None:
        labels = block.get("header_labels") or [column.get("id", "") for column in columns]
        row = table.add_row()
        if block.get("repeat_header", True):
            tables.repeat_header_row(row)

        for cell, label, column in zip(row.cells, labels, columns):
            tables.set_vertical_align(cell, vertical_align)
            paragraph = cell.paragraphs[0]
            style = column.get("header_style") or block.get("header_style")
            if style:
                paragraph.style = self.styles.paragraph(style)
            paragraph.add_run(render(label, context))

    def _fill_cell(self, cell, column: dict[str, Any], context: dict[str, Any]) -> None:
        text = render(column.get("value"), context)
        paragraph = cell.paragraphs[0]

        if column.get("style"):
            paragraph.style = self.styles.paragraph(column["style"])

        # Lien explicite déclaré dans le plan : il prime sur la détection
        # automatique lorsque son texte est présent dans la cellule.
        hyperlink = column.get("hyperlink") or {}
        link_text = render(hyperlink.get("text"), context)
        target = (
            self.links.bookmark_for(render(hyperlink.get("target"), context)) if link_text else ""
        )

        if (
            link_text
            and link_text in text
            and self.links.enabled
            and self.links.is_reachable(target)
            and evaluate(hyperlink.get("when"), context)
        ):
            before, _, after = text.partition(link_text)
            if before:
                self._write_rich_text(paragraph, before, context)
            self.links.add_hyperlink(paragraph, link_text, target)
            if after:
                self._write_rich_text(paragraph, after, context)
            return

        self._write_rich_text(paragraph, text, context, links=self._links_allowed(column, ""))

    def _write_loop(self, block: dict[str, Any], context: dict[str, Any], parent: str = "") -> None:
        """Répète une section et/ou des blocs sur chaque élément d'une collection."""
        item_name = block.get("item") or "item"
        section = block.get("section")

        for item in resolve_items(block.get("over"), context):
            item_context = {**context, item_name: item}
            if section:
                self._write_section(section, item_context, parent)
            for inner in block.get("blocks") or []:
                self._write_block(inner, item_context, parent)

    # ── Écriture du texte ─────────────────────────────────────────
    def _write_value(self, paragraph, value: Any, context: dict[str, Any], links: bool) -> None:
        """Écrit une valeur de liste : lien vers un signet, ou texte enrichi."""
        if isinstance(value, DocLink):
            target = self.links.bookmark_for(value.target)
            if links and self.links.enabled and self.links.is_reachable(target):
                self.links.add_hyperlink(paragraph, value.text, target)
                return
            self._write_rich_text(paragraph, value.text, context, links=links)
            return

        # Une chaîne provient des données (nom de colonne, étape Power Query) :
        # elle est écrite telle quelle, sans nouvelle substitution.
        text = value if isinstance(value, str) else render(value, context)
        self._write_rich_text(paragraph, text, context, links=links)

    def _write_rich_text(
        self,
        paragraph,
        text: str,
        context: dict[str, Any],
        bold: bool = False,
        links: bool = True,
    ) -> None:
        """Écrit un texte en gérant les retours à la ligne et les liens internes."""
        lines = str(text).split("\n")
        skip = self.links.self_bookmark(context) if links else None

        for index, line in enumerate(lines):
            segments = self.links.split(line, skip) if (links and line) else [(line, None)]
            for chunk, bookmark in segments:
                if bookmark:
                    self.links.add_hyperlink(paragraph, chunk, bookmark)
                elif chunk:
                    paragraph.add_run(chunk).bold = bold
            if index < len(lines) - 1:
                paragraph.add_run().add_break()

    def _links_allowed(self, node: dict[str, Any], style_name: str) -> bool:
        """Détermine si la détection automatique s'applique à ce contenu."""
        if not self.links.enabled or not self.links.auto.get("enabled", True):
            return False
        if node.get("links") is False or node.get("auto_links") is False:
            return False
        if not self.links.auto.get("in_code", True):
            return style_name != self.styles.paragraph("code")
        return True


def _numbering_mode(value: Any) -> str:
    """
    Mode de numérotation des figures : `auto`, `fixed` ou `none`.

    `auto` confie le numéro à un champ Word, qui le tient à jour lui-même —
    c'est le comportement voulu dans la quasi-totalité des cas. Les anciens
    plans écrivaient un booléen : `true` vaut `auto`, `false` vaut `none`.
    """
    if isinstance(value, bool) or value is None:
        return "auto" if value is not False else "none"
    mode = str(value).strip().lower()
    return mode if mode in ("auto", "fixed", "none") else "auto"


def _header_footer_parts(section) -> dict[str, tuple]:
    return {
        "header": (section.header, section.first_page_header, section.even_page_header),
        "footer": (section.footer, section.first_page_footer, section.even_page_footer),
    }


def _column_width(column: dict[str, Any]) -> float | None:
    """Largeur d'une colonne en centimètres (`width_cm`)."""
    width = column.get("width_cm")
    return float(width) if width else None


def _apply_column_widths(table, columns: list[dict[str, Any]]) -> None:
    for index, column in enumerate(columns):
        width = _column_width(column)
        if width is not None:
            for row in table.rows:
                row.cells[index].width = Cm(width)
