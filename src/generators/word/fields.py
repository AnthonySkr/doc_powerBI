"""
Champs Word et textes du template : table des matières, en-têtes, pieds de page.

Les numéros de page d'une table des matières dépendent de la mise en page :
seul Word peut les calculer. Le champ TOC est donc marqué « à recalculer »
(`w:dirty`), ce que Word applique à l'ouverture du document.
"""

import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Un champ de table des matières commence par le mot-clé TOC.
_TOC_FIELD = re.compile(r"^\s*TOC\b")

# Éléments de settings.xml devant suivre `w:updateFields` (ordre du schéma).
_SETTINGS_AFTER_UPDATE_FIELDS = (
    "w:hdrShapeDefaults",
    "w:footnotePr",
    "w:endnotePr",
    "w:compat",
    "w:docVars",
    "w:rsids",
    "w:mathPr",
    "w:attachedSchema",
    "w:themeFontLang",
    "w:clrSchemeMapping",
    "w:shapeDefaults",
    "w:decimalSymbol",
    "w:listSeparator",
)


# ─────────────────────────────────────────────────────────────
#  Table des matières
# ─────────────────────────────────────────────────────────────


def mark_toc_fields_dirty(body, levels: str = "") -> int:
    r"""
    Marque les champs TOC « à recalculer » et, si `levels` est renseigné,
    ajuste les niveaux de titres repris (`\o "1-3"`).

    Retourne le nombre de champs traités.
    """
    fields = 0

    # Champ complet : w:fldChar(begin) ... w:instrText ... w:fldChar(end)
    last_begin = None
    for element in body.iter(qn("w:fldChar"), qn("w:instrText")):
        if element.tag == qn("w:fldChar"):
            if element.get(qn("w:fldCharType")) == "begin":
                last_begin = element
            continue

        if not _TOC_FIELD.match(element.text or ""):
            continue

        if levels:
            element.text = set_toc_levels(element.text or "", levels)
        if last_begin is not None:
            last_begin.set(qn("w:dirty"), "true")
        fields += 1

    # Forme condensée : w:fldSimple w:instr="TOC ..."
    for element in body.iter(qn("w:fldSimple")):
        instruction = element.get(qn("w:instr")) or ""
        if not _TOC_FIELD.match(instruction):
            continue
        if levels:
            element.set(qn("w:instr"), set_toc_levels(instruction, levels))
        element.set(qn("w:dirty"), "true")
        fields += 1

    return fields


def set_toc_levels(instruction: str, levels: str) -> str:
    r"""
    Remplace les niveaux repris par le champ TOC (`\o "1-2"` → `\o "1-3"`).

    Le remplacement passe par une fonction : la chaîne contient un antislash,
    que `re.sub` interpréterait comme une séquence d'échappement.
    """
    return re.sub(r'\\o\s*"[^"]*"', lambda _: f'\\o "{levels}"', instruction)


def set_update_fields(settings) -> None:
    """
    Ajoute `<w:updateFields w:val="true"/>` : Word recalcule alors tous les
    champs du document à son ouverture. L'élément doit respecter l'ordre du
    schéma, sans quoi Word considère le fichier comme corrompu.
    """
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return

    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")

    for tag in _SETTINGS_AFTER_UPDATE_FIELDS:
        successor = settings.find(qn(tag))
        if successor is not None:
            successor.addprevious(element)
            return

    settings.append(element)


# ─────────────────────────────────────────────────────────────
#  Remplacement de texte dans un en-tête / pied de page
# ─────────────────────────────────────────────────────────────


def replace_in_part(part, placeholder: str, value: str) -> int:
    """Remplace un texte dans un en-tête ou un pied de page."""
    replaced = sum(
        replace_in_paragraph(paragraph, placeholder, value) for paragraph in part.paragraphs
    )
    for table in part.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced += replace_in_paragraph(paragraph, placeholder, value)
    return replaced


def replace_in_paragraph(paragraph, placeholder: str, value: str) -> int:
    """
    Remplace un texte dans un paragraphe sans toucher au reste de sa mise en
    forme : seuls les nœuds `w:t` sont réécrits, les tabulations et les images
    qui les entourent (logo, alignement à droite de l'en-tête) sont conservées.
    """
    nodes = paragraph._p.findall(".//" + qn("w:t"))
    if not nodes:
        return 0

    # Cas courant : le texte tient dans un seul nœud.
    replaced = 0
    for node in nodes:
        text = node.text or ""
        if placeholder in text:
            replaced += text.count(placeholder)
            _set_text(node, text.replace(placeholder, value))
    if replaced:
        return replaced

    # Le texte est réparti sur plusieurs nœuds : on le regroupe sur le premier.
    joined = "".join(node.text or "" for node in nodes)
    if placeholder not in joined:
        return 0

    _set_text(nodes[0], joined.replace(placeholder, value))
    for node in nodes[1:]:
        _set_text(node, "")
    return joined.count(placeholder)


def _set_text(node, text: str) -> None:
    node.text = text
    if text != text.strip():
        node.set(qn("xml:space"), "preserve")
