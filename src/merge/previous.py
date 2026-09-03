"""
Lecture du document généré précédemment.

Le document est ouvert, découpé en blocs ancrés, et **laissé ouvert** : la
fusion y puise non seulement du XML mais aussi les parties associées (une
capture collée par l'utilisateur vit dans une partie du .docx, pas dans son
paragraphe).

Le document n'est jamais modifié : il est lu, puis un document neuf est écrit.
"""

import os
from dataclasses import dataclass, field

from docx import Document

from src import console
from src.merge import blocks as block_parser
from src.merge import markers
from src.merge.blocks import Block

# États d'un élément vis-à-vis du document précédent.
NEW = "new"
CHANGED = "changed"
UNCHANGED = "unchanged"


@dataclass
class PreviousDocument:
    """Le document précédent, tel qu'il servira à la fusion."""

    path: str = ""
    document: object | None = None
    blocks: list[Block] = field(default_factory=list)
    fingerprints: dict[str, str] = field(default_factory=dict)

    @property
    def exists(self) -> bool:
        return bool(self.path)

    def status(self, element_id: str, fingerprint: str) -> str:
        """Compare un élément du rapport actuel à ce que contenait le document."""
        if not self.exists:
            return UNCHANGED
        previous = self.fingerprints.get(element_id)
        if previous is None:
            return NEW
        return CHANGED if fingerprint and previous != fingerprint else UNCHANGED

    def removed(self, written_ids: set[str]) -> list[str]:
        """
        Éléments présents dans le document précédent mais plus dans le rapport.

        Les ancres internes à la fusion — l'annexe des contenus non replacés —
        n'en font pas partie : elles ne décrivent rien du rapport, et se
        reconnaissent à leur préfixe.
        """
        return sorted(
            element_id
            for element_id in set(self.fingerprints) - written_ids
            if not element_id.startswith(markers.INTERNAL_PREFIX)
        )


def read(path: str) -> PreviousDocument:
    """Lit le document précédent, s'il existe et porte des marqueurs."""
    if not os.path.isfile(path):
        return PreviousDocument()

    try:
        document = Document(path)
    except Exception as e:  # noqa: BLE001
        console.warn(f"Document précédent illisible, il sera régénéré ({e})")
        return PreviousDocument()

    blocks = block_parser.parse(block_parser.body_nodes(document))
    anchored = block_parser.index(blocks)

    if not anchored:
        console.info("Document précédent sans marqueurs : il sera entièrement régénéré")
        return PreviousDocument()

    free = sum(len(block.free_nodes()) for block in anchored.values())
    console.done(
        f"version précédente relue : {len(anchored)} élément(s) repéré(s), "
        f"{free} paragraphe(s) et tableau(x) que vous avez écrits"
    )

    return PreviousDocument(
        path=path,
        document=document,
        blocks=blocks,
        fingerprints={key: block.fingerprint for key, block in anchored.items()},
    )
